import os
from datetime import datetime, date, timedelta
from docx import Document
from dateutil.relativedelta import relativedelta
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import (
    SCRIPT_DIR, load_contracts, generate_rent_plan,
    load_shops, load_payments,
    get_paid, update_paid, save_payments,
    total_rent, total_paid, arrears, remaining_rent
)
import utils

def _fuzzy_match(keyword, text):
    """模糊匹配：大小写不敏感子串匹配"""
    if not keyword:
        return True
    return keyword.lower() in str(text).lower()

def _render_inline_row(parent, c, fields, row, pady=(0, 6)):
    """将一组字段横向渲染为一行（标签：值 依次排列）"""
    row_frame = tk.Frame(parent)
    row_frame.grid(row=row, column=0, columnspan=4, sticky="w", pady=pady)
    for i, key in enumerate(fields):
        if i > 0:
            tk.Label(row_frame, text="　", font=("微软雅黑", 10),
                     width=4).pack(side="left")
        tk.Label(row_frame, text=f"{key}：", font=("微软雅黑", 10, "bold"),
                 fg="#333").pack(side="left")
        tk.Label(row_frame, text=str(c.get(key, "")), font=("微软雅黑", 10),
                 fg="#555").pack(side="left")

# ================== 主界面 ==================
class RentCollectionGUI:
    def __init__(self, parent):
        self.parent = parent
        self.contracts = []
        self._iid_to_idx = {}
        self.filter_keyword = tk.StringVar()
        self.refresh_data()
        self.create_widgets()

    def refresh_data(self):
        raw = load_contracts()
        self.contracts = [c for c in raw if c.get("合同状态") != "待生效"]

    def create_widgets(self):
        for w in self.parent.winfo_children():
            w.destroy()

        btn_frame = ttk.Frame(self.parent)
        btn_frame.pack(fill="x", padx=10)
        ttk.Button(btn_frame, text="收缴计划", command=self.show_plan).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="刷新", command=self.refresh_and_render).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="生成催缴函", command=self.generate_demand_letter).pack(side="left", padx=5)

        filter_frame = ttk.LabelFrame(self.parent, text="筛选条件", padding=6)
        filter_frame.pack(fill="x", padx=10, pady=5)

        kw_frame = ttk.Frame(filter_frame)
        kw_frame.grid(row=0, column=0, columnspan=8, sticky=tk.W, padx=5, pady=2)
        ttk.Label(kw_frame, text="关键词：").pack(side=tk.LEFT)
        ttk.Entry(kw_frame, textvariable=self.filter_keyword, width=14).pack(side=tk.LEFT, padx=2)
        ttk.Button(kw_frame, text="执行筛选", command=self.render_list).pack(side=tk.LEFT, padx=(6, 3))
        ttk.Button(kw_frame, text="清空条件", command=self.reset_filter).pack(side=tk.LEFT)

        ttk.Label(filter_frame, text="下次缴费剩余天数：").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        self.filter_min = tk.StringVar(value="0")
        self.filter_max = tk.StringVar(value="9999")
        ttk.Entry(filter_frame, textvariable=self.filter_min, width=5).grid(row=1, column=1)
        ttk.Label(filter_frame, text="~", width=2).grid(row=1, column=2)
        ttk.Entry(filter_frame, textvariable=self.filter_max, width=5).grid(row=1, column=3)

        self.filter_arrears = tk.BooleanVar()
        ttk.Checkbutton(filter_frame, text="已逾期", variable=self.filter_arrears).grid(row=1, column=4, padx=10)

        table_frame = ttk.Frame(self.parent)
        table_frame.pack(fill="both", expand=1, padx=10, pady=5)

        cols = ["序号", "合同号", "商户名称", "免租期(天)", "总租金", "总物业费", "总计预期收入",
                "总计已收租金", "总计已收物业费", "总计剩余收入",
                "下次缴费剩余天数", "租金逾期状态", "物业费逾期状态", "逾期金额"]
        self.tree = ttk.Treeview(table_frame, columns=cols, show="headings", selectmode="extended")
        for c in cols:
            self.tree.heading(c, text=c, command=lambda col=c: self._sort_by(col))
            self.tree.column(c, width=120)
        self.tree.column("序号", width=80)

        vs = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hs = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vs.grid(row=0, column=1, sticky="ns")
        hs.grid(row=1, column=0, sticky="ew")
        table_frame.rowconfigure(0, weight=1)
        table_frame.columnconfigure(0, weight=1)

        self.tree.tag_configure("even", background="#f7faff")
        self.tree.tag_configure("odd",  background="#ffffff")
        self.tree.tag_configure("summary", background="#f0f0f0", font=("微软雅黑", 10, "bold"))

        self._sort_col = "合同号"
        self._sort_rev = True
        self.render_list()

    def reset_filter(self):
        self.filter_min.set("0")
        self.filter_max.set("9999")
        self.filter_arrears.set(False)
        self.filter_keyword.set("")
        self.render_list()

    def _sort_by(self, col):
        if self._sort_col == col:
            self._sort_rev = not self._sort_rev
        else:
            self._sort_col = col
            self._sort_rev = True
        self.render_list()

    def render_list(self):
        try:
            self.tree.delete(*self.tree.get_children())
        except tk.TclError:
            return
        self._iid_to_idx.clear()
        fmin = int(self.filter_min.get() or 0)
        fmax = int(self.filter_max.get() or 9999)
        farr = self.filter_arrears.get()
        fkw = self.filter_keyword.get().strip()

        # 预加载缓存，避免每个合同重复加载
        from utils import load_shops, load_payments
        shops_cache = load_shops()
        payments_cache = load_payments()

        sum_total = 0.0        # 总租金累计
        sum_prop = 0.0          # 总物业费累计
        sum_paid = 0.0          # 总计已收租金累计
        sum_prop_paid = 0.0     # 总计已收物业费累计
        sum_expected = 0.0      # 总计预期收入累计
        sum_remaining = 0.0     # 总计剩余收入累计
        sum_arrears = 0.0
        from datetime import date as dt_date, datetime as dt_datetime
        today = dt_date.today()
        # 先收集所有行数据
        rows_data = []
        for idx, c in enumerate(self.contracts):
            # 关键词过滤（在生成租金计划前执行，节省性能）
            if fkw and not (_fuzzy_match(fkw, c.get("合同号", ""))
                            or _fuzzy_match(fkw, c.get("商户名称", ""))
                            or _fuzzy_match(fkw, c.get("关联铺位号", ""))):
                continue
            # 每个合同只生成一次租金计划，传入缓存
            plan = generate_rent_plan(c, _shops_cache=shops_cache, _payments_cache=payments_cache)
            # 从 plan 直接计算所有指标
            tr = round(sum(p["应缴金额(元)"] for p in plan), 2)
            pd = round(sum(p["已缴金额(元)"] for p in plan), 2)
            tp = round(sum(float(p.get("应缴物业费", 0)) for p in plan), 2)
            tpp = round(sum(float(p.get("已缴物业费", 0)) for p in plan), 2)
            te = round(tr + tp, 2)           # 总计预期收入
            trm = round(te - pd - tpp, 2)     # 总计剩余收入
            # 逾期金额（租金逾期 + 物业费逾期）
            ar = 0.0
            prop_ar = 0.0
            for p in plan:
                d = dt_datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
                if d < today:
                    ar += max(p["应缴金额(元)"] - p["已缴金额(元)"], 0)
                    prop_ar += max(p["应缴物业费"] - p["已缴物业费"], 0)
            ar = round(ar + prop_ar, 2)
            # 租金逾期状态 & 物业费逾期状态
            status = c.get("合同状态", "")
            # 未起租检查：当前日期 < 租赁开始日期
            start_str = c.get("租赁开始日期", "")
            if start_str:
                try:
                    start_dt_rent = dt_datetime.strptime(start_str, "%Y-%m-%d").date()
                    not_started = today < start_dt_rent
                except:
                    not_started = False
            else:
                not_started = False

            if status == "已到期":
                label, days, rent_status = "已到期", -1, "已到期"
            elif status == "已终止":
                label, days, rent_status = "已终止", -1, "已终止"
            elif not_started:
                label, days, rent_status = "未起租", -1, "未起租"
            else:
                unpaid = [p for p in plan if p["已缴金额(元)"] < p["应缴金额(元)"]]
                if not unpaid:
                    label, days, rent_status = "已结清", 9999, "已结清"
                else:
                    unpaid.sort(key=lambda x: x["支付时间"])
                    found = False
                    for p in unpaid:
                        d = dt_datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
                        if d >= today:
                            diff = (d - today).days
                            label, days, rent_status = str(diff), diff, "正常"
                            found = True
                            break
                    if not found:
                        label, days, rent_status = "已逾期", -1, "已逾期"
            # 物业费逾期状态
            if status in ("已到期", "已终止"):
                prop_status = status
            elif not_started:
                prop_status = "未起租"
            else:
                prop_unpaid = [
                    p for p in plan
                    if float(p.get("应缴物业费", 0)) > 0
                    and float(p.get("已缴物业费", 0)) < float(p.get("应缴物业费", 0))
                ]
                if not prop_unpaid:
                    prop_status = "已结清"
                else:
                    prop_unpaid.sort(key=lambda x: x["支付时间"])
                    prop_overdue = False
                    for p in prop_unpaid:
                        d = dt_datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
                        if d < today:
                            prop_overdue = True
                            break
                    prop_status = "已逾期" if prop_overdue else "正常"
            if farr and ar <= 0:
                continue
            if days != -1 and not (fmin <= days <= fmax):
                continue
            sum_total += tr
            sum_prop += tp
            sum_paid += pd
            sum_prop_paid += tpp
            sum_expected += te
            sum_remaining += trm
            sum_arrears += ar
            rows_data.append({
                "orig_idx": idx,
                "cno": c["合同号"],
                "name": c["商户名称"],
                "free_days": int(c.get("免租期(天)", 0) or 0),
                "tr": tr,
                "tp": tp,
                "te": te,
                "pd": pd,
                "tpp": tpp,
                "trm": trm,
                "label": label,
                "days": days,
                "ar": ar,
                "rent_status": rent_status,
                "prop_status": prop_status,
            })
        # 排序
        col_key_map = {
            "总租金": "tr", "总物业费": "tp", "总计预期收入": "te",
            "总计已收租金": "pd", "总计已收物业费": "tpp", "总计剩余收入": "trm",
            "逾期金额": "ar", "下次缴费剩余天数": "days",
            "合同号": "cno", "商户名称": "name", "免租期(天)": "free_days",
        }
        sort_key = col_key_map.get(self._sort_col, "cno")
        num_keys = {"tr", "tp", "te", "pd", "tpp", "trm", "ar", "days"}
        def _key(r):
            v = r.get(sort_key, "")
            if sort_key in num_keys:
                try: return float(v)
                except: return 0.0
            return str(v)
        rows_data.sort(key=_key, reverse=self._sort_rev)
        # 插入排序后的行
        for i, rd in enumerate(rows_data):
            tag = "even" if i % 2 == 0 else "odd"
            iid = self.tree.insert("", "end", values=[
                i + 1, rd["cno"], rd["name"], rd["free_days"], rd["tr"], rd["tp"],
                rd["te"], rd["pd"], rd["tpp"], rd["trm"],
                rd["label"], rd["rent_status"], rd["prop_status"], rd["ar"]
            ], tags=(tag,))
            self._iid_to_idx[iid] = rd["orig_idx"]
        self.tree.insert("", "end", values=[
            "合计", "", "", "",
            round(sum_total, 2),
            round(sum_prop, 2),
            round(sum_expected, 2),
            round(sum_paid, 2),
            round(sum_prop_paid, 2),
            round(sum_remaining, 2),
            "", "", "",
            round(sum_arrears, 2),
        ], tags=("summary",))

    def get_selected(self):
        selected = []
        for iid in self.tree.selection():
            item = self.tree.item(iid)
            if "summary" in item["tags"]:
                continue
            idx = self._iid_to_idx.get(iid)
            if idx is not None and idx < len(self.contracts):
                selected.append(self.contracts[idx])
        return selected

    def refresh_and_render(self):
        self.refresh_data()
        self.render_list()

    def show_detail(self):
        s = self.get_selected()
        if not s:
            messagebox.showwarning("提示", "请勾选合同")
            return
        win = tk.Toplevel(self.parent)
        win.title("合同详情")
        win.geometry("700x500")
        txt = tk.Text(win, wrap="word")
        scr = ttk.Scrollbar(win, command=txt.yview)
        txt.configure(yscrollcommand=scr.set)
        txt.pack(side="left", fill="both", expand=1)
        scr.pack(side="right", fill="y")
        for c in s:
            txt.insert("end", f"=== {c['合同号']} | {c['商户名称']} ===\n")
            for k, v in c.items():
                if k != "自动租金计划":
                    txt.insert("end", f"{k}：{v}\n")
            txt.insert("end", "\n")
        txt.config(state="disabled")

    def show_plan(self):
        s = self.get_selected()
        if not s:
            messagebox.showwarning("提示", "请选择合同")
            return
        if len(s) > 1:
            messagebox.showwarning("提示", "一次只能查看一个合同")
            return
        c = s[0]
        cno = c["合同号"]
        plan = generate_rent_plan(c)
        today = date.today()
        # 直接从 plan 计算汇总，避免重复生成租金计划
        tr = round(sum(p["应缴金额(元)"] for p in plan), 2)
        pd = round(sum(p["已缴金额(元)"] for p in plan), 2)
        rm = round(tr - pd, 2)
        ar = 0.0
        for p in plan:
            d = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
            if d < today:
                ar += max(p["应缴金额(元)"] - p["已缴金额(元)"], 0)
                ar += max(p["应缴物业费"] - p["已缴物业费"], 0)
        ar = round(ar, 2)
        win = tk.Toplevel(self.parent)
        win.title(f"收缴计划 - {cno}")
        win.geometry("1150x750")

        # ===== 合同详情（三层布局）=====
        detail_frame = tk.Frame(win, bd=1, relief="solid")
        detail_frame.pack(fill="x", padx=10, pady=(8, 0))

        inner = tk.Frame(detail_frame)
        inner.pack(fill="x", padx=12, pady=(8, 8))

        # 未起租判断
        start_str = c.get("租赁开始日期", "")
        try:
            _start_dt = datetime.strptime(start_str, "%Y-%m-%d").date()
            _not_started = today < _start_dt
        except:
            _not_started = False

        # ── 第1层：标题行（核心标识，大字加粗）──
        title_parts = [
            str(c.get("合同号", "")),
            str(c.get("商户名称", "")),
            str(c.get("经营业态", "")),
            str(c.get("所属项目", "")),
            str(c.get("关联铺位号", "")),
        ]
        title_text = "　".join(title_parts)
        tk.Label(inner, text=title_text, font=("微软雅黑", 13, "bold"),
                 fg="#1a5276", anchor="w").grid(row=0, column=0, columnspan=4,
                 sticky="w", pady=(0, 2))

        # ── 第2层：联系人栏 ──
        contact_text = f"联系人：{c.get('联系人', '')}　　联系电话：{c.get('联系电话', '')}"
        tk.Label(inner, text=contact_text, font=("微软雅黑", 10),
                 fg="#333").grid(row=1, column=0, columnspan=4,
                 sticky="w", pady=(0, 4))

        # ── 分隔线 ──
        tk.Frame(inner, height=1, bg="#ddd").grid(row=2, column=0,
                 columnspan=4, sticky="ew", pady=(0, 8))

        # ── 第3层：主体（两行横向排列）──
        # 第1行（上）：费用/押金类
        top_row_fields = [
            "租金模式",
            "保底租金(元/㎡/天)",
            "提成租金扣点(%)",
            "物业服务费单价（元/㎡/天）",
            "押金",
            "押金支付状态",
        ]
        _render_inline_row(inner, c, top_row_fields, row=3, pady=(0, 6))

        # 第2行（下）：日期/周期类
        bottom_row_fields = [
            "签约日期",
            "租赁开始日期",
            "租赁结束日期",
            "免租期(天)",
            "终止日期",
        ]
        _render_inline_row(inner, c, bottom_row_fields, row=4, pady=(0, 0))

        # ── 分隔线（主体 → 底行）──
        tk.Frame(inner, height=1, bg="#ddd").grid(row=5, column=0,
                 columnspan=4, sticky="ew", pady=(8, 8))

        # ── 底行：合同状态（值变色）+ 剩余租期（全黑），左对齐 ──
        status_val = c.get("合同状态", "")
        status_color_map = {
            "执行中": "#27ae60",
            "已到期": "#c0392b",
            "已终止": "#7f8c8d",
            "未起租": "#2980b9",
        }
        status_color = status_color_map.get(status_val, "#333")
        remain_display = "未起租" if _not_started else str(c.get("剩余租期(天)", ""))

        bottom_frame = tk.Frame(inner)
        bottom_frame.grid(row=6, column=0, columnspan=4, sticky="w", pady=(8, 0))
        tk.Label(bottom_frame, text="合同状态：", font=("微软雅黑", 10, "bold"),
                 fg="#333").pack(side="left")
        tk.Label(bottom_frame, text=status_val, font=("微软雅黑", 10, "bold"),
                 fg=status_color).pack(side="left")
        tk.Label(bottom_frame, text=f"　　剩余租期(天)：{remain_display}",
                 font=("微软雅黑", 10, "bold"), fg="#333").pack(side="left")

        # ── 合同备注（独立一行）──
        remark_text = f"合同备注：{c.get('合同备注', '')}"
        tk.Label(inner, text=remark_text, font=("微软雅黑", 10),
                 fg="#555").grid(row=7, column=0, columnspan=4,
                 sticky="w", pady=(4, 0))

        # ===== 收缴计划表格 =====
        fr = ttk.Frame(win)
        fr.pack(fill="both", expand=1, padx=10, pady=5)
        cols = ["序号", "支付时间", "应缴租金", "应缴物业费", "本期应缴",
                "已缴租金（双击录入）", "已缴物业费（双击录入）"]
        tree = ttk.Treeview(fr, columns=cols, show="headings")
        col_widths = {"序号": 50, "支付时间": 110, "应缴租金": 120,
                      "应缴物业费": 120, "本期应缴": 120,
                      "已缴租金（双击录入）": 160, "已缴物业费（双击录入）": 160}
        for col in cols:
            tree.heading(col, text=col)
            tree.column(col, width=col_widths.get(col, 150))
        vs = ttk.Scrollbar(fr, orient="vertical", command=tree.yview)
        hs = ttk.Scrollbar(fr, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
        tree.grid(row=0, column=0, sticky="nsew")
        vs.grid(row=0, column=1, sticky="ns")
        hs.grid(row=1, column=0, sticky="ew")
        fr.rowconfigure(0, weight=1)
        fr.columnconfigure(0, weight=1)
        tree.tag_configure("even", background="#f7faff")
        tree.tag_configure("odd",  background="#ffffff")
        tree.tag_configure("overdue", background="#ffeeee", foreground="red")
        tree.tag_configure("summary", background="#f0f0f0", font=("微软雅黑", 10, "bold"))
        total_should_val = 0.0
        total_paid_val = 0.0
        total_prop_val = 0.0
        total_prop_paid_val = 0.0
        for idx, p in enumerate(plan, 1):
            dt = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
            s_val = p["应缴金额(元)"]
            pd_val = p["已缴金额(元)"]
            prop_val = p["应缴物业费"]
            prop_pd_val = p["已缴物业费"]
            total_should_val += s_val
            total_paid_val += pd_val
            total_prop_val += prop_val
            total_prop_paid_val += prop_pd_val
            parity = "even" if (idx - 1) % 2 == 0 else "odd"
            if dt < today and (pd_val < s_val or prop_pd_val < prop_val):
                tag = ("overdue", parity)
            else:
                tag = (parity,)
            tree.insert("", "end", values=[idx, p["支付时间"], s_val, prop_val, round(s_val + prop_val, 2), pd_val, prop_pd_val], tags=tag)
        tree.insert("", "end", values=["汇总", "", round(total_should_val, 2), round(total_prop_val, 2),
                                         round(total_should_val + total_prop_val, 2),
                                         round(total_paid_val, 2), round(total_prop_paid_val, 2)], tags=("summary",))

        def edit(evt):
            row = tree.identify_row(evt.y)
            col = tree.identify_column(evt.x)
            if not row or "summary" in tree.item(row)["tags"]:
                return
            col_idx = int(col.replace("#", ""))
            if col_idx == 6:  # 已缴租金列
                val = tree.item(row)["values"][5]
                value_index = 5
            elif col_idx == 7:  # 已缴物业费列
                val = tree.item(row)["values"][6]
                value_index = 6
            else:
                return
            x, y, w, h = tree.bbox(row, col)
            ent = ttk.Entry(tree)
            ent.place(x=x, y=y, width=w, height=h)
            ent.insert(0, val)
            ent.focus()

            def save(evt2):
                try:
                    new_val = round(float(ent.get()), 2)
                except:
                    new_val = 0.0
                v = list(tree.item(row)["values"])
                v[value_index] = new_val
                tree.item(row, values=v)
                ent.destroy()
            ent.bind("<Return>", save)
            ent.bind("<FocusOut>", save)
        tree.bind("<Double-1>", edit)

        def save_all():
            from utils import update_property_fee_paid
            for row in tree.get_children():
                if "summary" in tree.item(row)["tags"]:
                    continue
                v = tree.item(row)["values"]
                update_paid(cno, v[1], v[5])
                update_property_fee_paid(cno, v[1], v[6])
            utils.log_operation("修改", "租金", f"编辑收缴计划 {cno}（{c['商户名称']}）", cno)
            messagebox.showinfo("成功", "已保存")
            win.destroy()
            self.refresh_and_render()
        ttk.Button(win, text="💾 保存", command=save_all).pack(pady=5)

    # ========== 生成催缴函：一次选文件夹、多份批量 ==========
    def generate_demand_letter(self):
        selected = self.get_selected()
        if not selected:
            messagebox.showwarning("提示", "请先选择合同")
            return

        save_dir = filedialog.askdirectory(title="选择催缴函保存目录")
        if not save_dir:
            return

        today = date.today()
        ymd = today.strftime("%Y年%m月%d日")
        success = 0
        fail_list = []

        for c in selected:
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                doc = Document()

                section = doc.sections[0]
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)

                style = doc.styles['Normal']
                style.font.name = '微软雅黑'
                style.font.size = Pt(14)
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                def _set_font(run, size=None, bold=None):
                    run.font.name = '微软雅黑'
                    if size:
                        run.font.size = size
                    if bold is not None:
                        run.bold = bold
                    rPr = run._element.get_or_add_rPr()
                    existing = rPr.findall(qn('w:rFonts'))
                    found = False
                    for e in existing:
                        if e.get(qn('w:eastAsia')) == '微软雅黑':
                            found = True
                            break
                    if not found:
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        rPr.insert(0, rFonts)

                def _set_indent(para, chars=2):
                    pPr = para._element.get_or_add_pPr()
                    for old in pPr.findall(qn('w:ind')):
                        pPr.remove(old)
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:firstLine'), str(chars * 280))
                    ind.set(qn('w:firstLineChars'), str(chars * 100))
                    pPr.append(ind)

                def _set_para_spacing(para):
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.0

                def _add_body(text, indent_chars=2):
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    _set_font(run, size=Pt(14))
                    _set_para_spacing(p)
                    if indent_chars > 0:
                        _set_indent(p, indent_chars)
                    return p

                def _add_empty(indent_chars=0):
                    p = doc.add_paragraph()
                    _set_para_spacing(p)
                    if indent_chars > 0:
                        _set_indent(p, indent_chars)
                    return p

                p = doc.add_paragraph()
                run = p.add_run("租金催缴函")
                _set_font(run, size=Pt(16), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                _add_empty(indent_chars=2)

                p = doc.add_paragraph()
                run = p.add_run(f"尊敬的 {c['商户名称']}：")
                _set_font(run, size=Pt(14))
                _set_para_spacing(p)

                _add_empty(indent_chars=2)

                # ── 欠缴信息：分别计算租金和物业费 ──
                plan = generate_rent_plan(c)
                rent_rows = []
                prop_rows = []
                rent_sub = 0.0
                prop_sub = 0.0
                for p_data in plan:
                    dt = datetime.strptime(p_data["支付时间"], "%Y-%m-%d").date()
                    if dt < today:
                        rd = round(max(p_data["应缴金额(元)"] - p_data["已缴金额(元)"], 0), 2)
                        ps = round(float(p_data.get("应缴物业费", 0) or 0), 2)
                        pp = round(float(p_data.get("已缴物业费", 0) or 0), 2)
                        pd = round(max(ps - pp, 0), 2)
                        if rd > 0:
                            rent_rows.append((p_data["支付时间"], round(p_data["应缴金额(元)"], 2), round(p_data["已缴金额(元)"], 2), rd))
                            rent_sub += rd
                        if pd > 0:
                            prop_rows.append((p_data["支付时间"], ps, pp, pd))
                            prop_sub += pd
                rent_sub = round(rent_sub, 2)
                prop_sub = round(prop_sub, 2)
                arrears_total = round(rent_sub + prop_sub, 2)

                _add_body(f"基于贵我双方签订的{c['合同号']}合同，您已欠缴我司费用合计：")

                # 总金额突出显示（单列表格框）
                _atbl = doc.add_table(rows=1, cols=1)
                _atbl.style = 'Table Grid'
                _ac = _atbl.rows[0].cells[0]
                _ac.text = ""
                _ap = _ac.paragraphs[0]
                _ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _ar = _ap.add_run(f"\uffe5{arrears_total}")
                _set_font(_ar, size=Pt(24), bold=True)
                _set_para_spacing(_ap)

                _add_empty()

                p = doc.add_paragraph()
                run = p.add_run(f"合同编号：{c['合同号']}")
                _set_font(run, size=Pt(10.5))
                _set_para_spacing(p)

                p = doc.add_paragraph()
                run = p.add_run(f"租赁起止日期：{c['租赁开始日期']} 至 {c['租赁结束日期']}")
                _set_font(run, size=Pt(10.5))
                _set_para_spacing(p)

                _add_empty()

                # ── 逾期明细表格辅助函数 ──
                def _add_arrears_table(title, headers, data_rows, subtotal_val, col_widths):
                    p = doc.add_paragraph()
                    run = p.add_run(title)
                    _set_font(run, size=Pt(11), bold=True)
                    _set_para_spacing(p)

                    ncols = len(headers)
                    tbl = doc.add_table(rows=1, cols=ncols)
                    tbl.style = 'Table Grid'
                    tw = sum(col_widths)
                    _el = tbl._tbl
                    _pr = _el.find(qn('w:tblPr'))
                    if _pr is None:
                        _pr = OxmlElement('w:tblPr')
                        _el.insert(0, _pr)
                    _tw = _pr.find(qn('w:tblW'))
                    if _tw is None:
                        _tw = OxmlElement('w:tblW')
                        _pr.insert(0, _tw)
                    _tw.set(qn('w:w'), str(tw))
                    _tw.set(qn('w:type'), 'dxa')
                    _gr = _el.find(qn('w:tblGrid'))
                    if _gr is None:
                        _gr = OxmlElement('w:tblGrid')
                        _el.insert(1, _gr)
                    for _gc in _gr.findall(qn('w:gridCol')):
                        _gr.remove(_gc)
                    for w_val in col_widths:
                        gc = OxmlElement('w:gridCol')
                        gc.set(qn('w:w'), str(w_val))
                        _gr.append(gc)

                    for i, text in enumerate(headers):
                        c = tbl.rows[0].cells[i]
                        c.text = ""
                        hp = c.paragraphs[0]
                        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        hr = hp.add_run(text)
                        _set_font(hr, size=Pt(10.5))
                        _set_para_spacing(hp)

                    for seq, rd in enumerate(data_rows, 1):
                        r = tbl.add_row().cells
                        vals = [str(seq), str(rd[0]), str(rd[1]), str(rd[2]), str(rd[3])]
                        for i, val in enumerate(vals):
                            r[i].text = ""
                            vp = r[i].paragraphs[0]
                            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            vr = vp.add_run(val)
                            _set_font(vr, size=Pt(10.5))
                            _set_para_spacing(vp)

                    # 小计行
                    sub_cells = tbl.add_row().cells
                    sub_cells[0].text = ""
                    sp0 = sub_cells[0].paragraphs[0]
                    sp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sr0 = sp0.add_run("小计")
                    _set_font(sr0, size=Pt(10.5), bold=True)
                    _set_para_spacing(sp0)
                    sub_cells[1].text = ""
                    sub_cells[2].text = ""
                    sub_cells[3].text = ""
                    sp3 = sub_cells[3].paragraphs[0]
                    sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sr3 = sp3.add_run("\u2014")
                    _set_font(sr3, size=Pt(10.5))
                    _set_para_spacing(sp3)
                    sub_cells[4].text = ""
                    sp4 = sub_cells[4].paragraphs[0]
                    sp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sr4 = sp4.add_run(str(subtotal_val))
                    _set_font(sr4, size=Pt(10.5), bold=True)
                    _set_para_spacing(sp4)

                    for row_obj in tbl.rows:
                        for i, cell in enumerate(row_obj.cells):
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcW = tcPr.find(qn('w:tcW'))
                            if tcW is None:
                                tcW = OxmlElement('w:tcW')
                                tcPr.insert(0, tcW)
                            tcW.set(qn('w:w'), str(col_widths[i]))
                            tcW.set(qn('w:type'), 'dxa')

                    _add_empty()

                # 5列宽度：序号/日期/应缴/已缴/欠付
                CW5 = [1000, 1600, 2200, 2100, 2135]

                # 租金欠缴明细表
                if rent_rows:
                    _add_arrears_table("租金欠缴明细：",
                                       ["序号", "应支付日期", "应缴金额（元）", "已缴金额（元）", "欠付金额（元）"],
                                       rent_rows, rent_sub, CW5)

                # 物业费欠缴明细表
                if prop_rows:
                    _add_arrears_table("物业费欠缴明细：",
                                       ["序号", "应支付日期", "应缴物业费（元）", "已缴物业费（元）", "欠付物业费（元）"],
                                       prop_rows, prop_sub, CW5)

                _add_body("请于收函之日起3日内将欠缴费用支付至以下账户，否则我们将采取进一步措施。")

                _add_empty()

                _add_body("开户名称：XXX公司", indent_chars=4)
                _add_body("开户行：XXX银行", indent_chars=4)
                _add_body("账号：11111111111", indent_chars=4)

                _add_empty(indent_chars=2)

                _add_body("顺颂商祺。")

                _add_empty()

                p_company = doc.add_paragraph()
                run = p_company.add_run("xxx公司")
                _set_font(run, size=Pt(14))
                _set_para_spacing(p_company)
                p_company.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                p_date = doc.add_paragraph()
                run = p_date.add_run(ymd)
                _set_font(run, size=Pt(14))
                _set_para_spacing(p_date)
                p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                filename = f"租金催缴函_{c['商户名称']}_{c['合同号']}.docx"
                out_path = os.path.join(save_dir, filename)
                counter = 1
                while os.path.exists(out_path):
                    filename = f"租金催缴函_{c['商户名称']}_{c['合同号']}_{counter}.docx"
                    out_path = os.path.join(save_dir, filename)
                    counter += 1

                doc.save(out_path)
                success += 1

            except Exception as e:
                fail_list.append(f"{c['合同号']}：{str(e)}")

        msg = f"批量生成完成！\n成功：{success} 份"
        if fail_list:
            msg += f"\n失败：{len(fail_list)} 份\n" + "\n".join(fail_list)
        messagebox.showinfo("生成结果", msg)

if __name__ == "__main__":
    root = tk.Tk()
    root.title("租金收缴管理系统")
    root.geometry("1200x800")
    root.state("zoomed")
    app = RentCollectionGUI(root)
    root.mainloop()
