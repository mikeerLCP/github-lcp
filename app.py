"""
app.py —— 丰盛云商商管系统（Web版）应用入口
"""
import json, io, zipfile, os
from datetime import datetime, date, timedelta
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file
from config import SECRET_KEY
import auth
import db
import utils

app = Flask(__name__)
app.secret_key = SECRET_KEY

# ===================== 登录 =====================
@app.route("/")
def index():
    """首页重定向到登录"""
    if auth.is_logged_in():
        return redirect(url_for("admin_dashboard"))
    return redirect(url_for("login_page"))

@app.route("/login", methods=["GET", "POST"])
def login_page():
    """登录页"""
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        if not username or not password:
            flash("请输入用户名和密码", "warning")
            return render_template("login.html")
        user = auth.login_user(username, password)
        if user:
            db.log_operation(user["用户名"], user["角色"], "登录", "系统", "登录成功")
            return redirect(url_for("admin_dashboard"))
        else:
            flash("用户名或密码错误", "danger")
    return render_template("login.html")

@app.route("/logout")
def logout():
    """退出登录"""
    user = auth.get_current_user()
    if user:
        db.log_operation(user["用户名"], user["角色"], "登出", "系统", "退出登录")
    auth.logout_user()
    flash("已退出登录", "info")
    return redirect(url_for("login_page"))

# ===================== 管理端 - 仪表板 =====================
@app.route("/admin")
@auth.login_required
def admin_dashboard():
    """管理端仪表板"""
    from datetime import date
    user = auth.get_current_user()
    project_filter = auth.get_project_filter()
    year = request.args.get("year", "")
    year = int(year) if year.isdigit() else date.today().year
    stats = db.get_dashboard_stats(project_filter, year)
    biz_data = db.get_business_dashboard_api(project_filter)
    return render_template("admin/dashboard.html", user=user, stats=stats, biz_data=biz_data, sel_year=year)


@app.route("/api/biz-dashboard")
@auth.login_required
def api_biz_dashboard():
    """AJAX：经营数据看板数据"""
    project_filter = auth.get_project_filter()
    date_str = request.args.get("date", "").strip()
    biz_project = request.args.get("project", "").strip()
    return jsonify(db.get_business_dashboard_api(project_filter, date_str, biz_project))

# ===================== 管理端 - 商铺管理 =====================
@app.route("/admin/shops")
@auth.admin_required
def admin_shops():
    """商铺列表"""
    user = auth.get_current_user()
    project_filter = auth.get_project_filter()
    shops = db.load_shops(project_filter)
    from config import PROJECT_OPTIONS, BUSINESS_TYPE, SHOP_STATUS
    return render_template("admin/shops.html", user=user, shops=shops,
                           projects=PROJECT_OPTIONS, biz_types=BUSINESS_TYPE,
                           shop_statuses=SHOP_STATUS)

@app.route("/admin/shops/add", methods=["POST"])
@auth.admin_required
def admin_shops_add():
    """新增/编辑商铺"""
    user = auth.get_current_user()
    shop_data = {
        "铺位号": request.form.get("铺位号", "").strip(),
        "所属项目": request.form.get("所属项目", "").strip(),
        "位置": request.form.get("位置", "").strip(),
        "铺位状态": request.form.get("铺位状态", "空置"),
        "空间类型": request.form.get("空间类型", "").strip(),
        "上下水": request.form.get("上下水", "").strip(),
        "电力功率上限(kW)": request.form.get("电力功率上限", "").strip(),
        "装修情况": request.form.get("装修情况", "").strip(),
        "租金报价(元/㎡/天)": request.form.get("租金报价", "").strip(),
        "改造条件": request.form.get("改造条件", "").strip(),
        "户型图路径": request.form.get("户型图路径", "").strip(),
        "现状照片路径": request.form.get("现状照片路径", "").strip(),
        "建筑面积(㎡)": request.form.get("建筑面积", "").strip(),
        "使用面积(㎡)": request.form.get("使用面积", "").strip(),
        "基准租金(元/㎡/天)": request.form.get("基准租金", "").strip(),
        "备注": request.form.get("备注", "").strip(),
    }
    # ── 自动填充计租面积 ──
    st = shop_data.get("空间类型", "")
    ba = shop_data.get("建筑面积(㎡)", "")
    ua = shop_data.get("使用面积(㎡)", "")
    if st == "商铺":
        shop_data["计租面积(㎡)"] = ba
    elif st == "室外场地":
        shop_data["计租面积(㎡)"] = ua
    else:
        shop_data["计租面积(㎡)"] = ""
    # ── 后端校验 ──
    if not shop_data["铺位号"]:
        flash("铺位号不能为空", "danger")
        return redirect(url_for("admin_shops"))
    if not shop_data["所属项目"]:
        flash("请选择所属项目", "danger")
        return redirect(url_for("admin_shops"))
    # 数值格式校验
    for key, label in [("建筑面积(㎡)", "建筑面积"), ("使用面积(㎡)", "使用面积"), ("基准租金(元/㎡/天)", "基准租金")]:
        val = shop_data[key]
        if val:
            try:
                f = float(val)
                if f <= 0:
                    flash(f"{label}需大于0", "danger")
                    return redirect(url_for("admin_shops"))
            except ValueError:
                flash(f"{label}格式不正确", "danger")
                return redirect(url_for("admin_shops"))
    is_new = not db.check_shop_no_exists(shop_data["铺位号"])
    db.save_shops([shop_data])
    op = "新增" if is_new else "修改"
    db.log_operation(user["用户名"], user["角色"], op, "空间",
                     f"{op}空间 {shop_data['铺位号']}", shop_data["铺位号"])
    flash(f"空间 {shop_data['铺位号']} 保存成功", "success")
    return redirect(url_for("admin_shops"))

@app.route("/admin/shops/delete/<shop_no>")
@auth.admin_required
def admin_shops_delete(shop_no):
    """删除商铺"""
    user = auth.get_current_user()
    # 检查关联合同
    contracts = db.get_contracts_by_shop(shop_no)
    if contracts:
        names = ", ".join([c["商户名称"] for c in contracts])
        flash(f"无法删除：铺位 {shop_no} 关联了合同（{names}），请先处理关联合同", "danger")
        return redirect(url_for("admin_shops"))
    db.delete_shop(shop_no)
    db.log_operation(user["用户名"], user["角色"], "删除", "空间",
                     f"删除空间 {shop_no}", shop_no)
    flash(f"空间 {shop_no} 已删除", "success")
    return redirect(url_for("admin_shops"))

@app.route("/admin/shops/batch_delete", methods=["POST"])
@auth.admin_required
def admin_shops_batch_delete():
    """批量删除商铺"""
    user = auth.get_current_user()
    shop_nos = request.form.getlist("shop_nos")
    if not shop_nos:
        flash("未选择任何空间", "danger")
        return redirect(url_for("admin_shops"))
    deleted = []
    skipped = []
    for shop_no in shop_nos:
        contracts = db.get_contracts_by_shop(shop_no)
        if contracts:
            names = ", ".join([c["商户名称"] for c in contracts])
            skipped.append(f"{shop_no}（关联合同：{names}）")
        else:
            db.delete_shop(shop_no)
            db.log_operation(user["用户名"], user["角色"], "批量删除", "空间",
                             f"批量删除空间 {shop_no}", shop_no)
            deleted.append(shop_no)
    if deleted:
        flash(f"已删除 {len(deleted)} 个空间：{', '.join(deleted)}", "success")
    if skipped:
        flash(f"跳过 {len(skipped)} 个（有关联合同）：{'; '.join(skipped)}", "danger")
    return redirect(url_for("admin_shops"))

@app.route("/api/shop-precheck")
@auth.admin_required
def api_shop_precheck():
    """AJAX：检查铺位号是否已存在"""
    no = request.args.get("no", "").strip()
    if not no:
        return jsonify({"exists": False})
    exists = db.check_shop_no_exists(no)
    return jsonify({"exists": exists})

# ===================== 管理端 - 合同管理 =====================

def _sync_shop_status():
    """根据合同状态自动同步铺位状态（对齐原型 sync_shop_status）"""
    contracts = db.load_contracts()
    shops = db.load_shops()
    active_shops = set()
    for c in contracts:
        sno = (c.get("关联铺位号") or "").strip()
        status = c.get("合同状态", "")
        if sno and status in ("待生效", "执行中"):
            active_shops.add(sno)
    all_contracted = set()
    for c in contracts:
        sno = (c.get("关联铺位号") or "").strip()
        if sno:
            all_contracted.add(sno)
    modified = False
    for shop in shops:
        sno = shop.get("铺位号", "")
        if sno in active_shops:
            if shop.get("铺位状态") != "已出租":
                shop["铺位状态"] = "已出租"
                modified = True
        elif sno in all_contracted:
            if shop.get("铺位状态") != "空置":
                shop["铺位状态"] = "空置"
                modified = True
    if modified:
        db.save_shops(shops)

def _auto_expire_contracts():
    """遍历所有合同，租赁结束日期已过期且状态为「执行中」的自动改为「已到期」（对齐原型 auto_expire_contracts）"""
    contracts = db.load_contracts()
    today = date.today()
    modified = False
    for c in contracts:
        status = c.get("合同状态", "")
        end_str = c.get("租赁结束日期", "")
        if status == "执行中" and end_str:
            try:
                end_dt = datetime.strptime(end_str, "%Y-%m-%d").date()
                if end_dt < today:
                    c["合同状态"] = "已到期"
                    c["剩余租期(天)"] = "0"
                    modified = True
            except Exception:
                pass
    if modified:
        db.save_contracts(contracts)
    return modified

def _auto_activate_contracts():
    """将签约日期≤今天的待生效合同翻转为执行中"""
    contracts = db.load_contracts()
    today = date.today()
    modified = False
    for c in contracts:
        status = c.get("合同状态", "")
        sign_str = c.get("签约日期", "")
        if status == "待生效" and sign_str:
            try:
                sign_dt = datetime.strptime(sign_str, "%Y-%m-%d").date()
                if sign_dt <= today:
                    c["合同状态"] = "执行中"
                    modified = True
            except Exception:
                pass
    if modified:
        db.save_contracts(contracts)
    return modified

@app.route("/admin/contracts")
@auth.admin_required
def admin_contracts():
    """合同列表"""
    # 先自动检查到期合同、自动激活待生效合同，再同步铺位状态
    try:
        _auto_expire_contracts()
        _auto_activate_contracts()
        _sync_shop_status()
    except Exception as e:
        print(f"[合同自动检查] {e}")
    user = auth.get_current_user()
    project_filter = auth.get_project_filter()
    contracts = db.load_contracts(project_filter)
    shops = {s["铺位号"]: s for s in db.load_shops(project_filter)}
    opps = db.load_opportunities(project_filter)
    from config import PROJECT_OPTIONS, BUSINESS_TYPE
    return render_template("admin/contracts.html", user=user, contracts=contracts,
                           shops=shops, projects=PROJECT_OPTIONS, biz_types=BUSINESS_TYPE,
                           opps=opps, today=date.today().isoformat())

@app.route("/admin/contracts/add", methods=["POST"])
@auth.admin_required
def admin_contracts_add():
    """新增/编辑合同（对齐原型：自动计算剩余租期、同步商机阶段、同步铺位状态）"""
    user = auth.get_current_user()
    is_edit = request.form.get("is_edit", "0") == "1"
    contract_data = {
        "合同号": request.form.get("合同号", "").strip(),
        "商户名称": request.form.get("商户名称", "").strip(),
        "经营业态": request.form.get("经营业态", "").strip(),
        "所属项目": request.form.get("所属项目", "").strip(),
        "关联铺位号": request.form.get("关联铺位号", "").strip(),
        "保底租金(元/㎡/天)": request.form.get("保底租金", "0"),
        "提成租金扣点(%)": request.form.get("提成扣点", "0"),
        "签约日期": request.form.get("签约日期", ""),
        "租赁开始日期": request.form.get("租赁开始日期", ""),
        "租赁结束日期": request.form.get("租赁结束日期", ""),
        "终止日期": request.form.get("终止日期", ""),
        "免租期(天)": request.form.get("免租期", "0"),
        "押金": request.form.get("押金", "0"),
        "押金支付状态": request.form.get("押金支付状态", "").strip(),
        "意向金抵扣押金": request.form.get("意向金抵扣押金", "0").strip(),
        "已付补缴押金": request.form.get("已付补缴押金", "0").strip(),
        "支付周期": request.form.get("支付周期", "").strip(),
        "合同状态": request.form.get("合同状态", "").strip(),
        "联系电话": request.form.get("联系电话", "").strip(),
        "联系人": request.form.get("联系人", ""),
        "备注": request.form.get("备注", ""),
        "签约主体": request.form.get("签约主体", ""),
        "租金模式": request.form.get("租金模式", "保底").strip(),
        "物业服务费单价（元/㎡/天）": request.form.get("物业费单价", "0"),
        "终止原因": request.form.get("终止原因", "").strip(),
        "前序合同号": request.form.get("前序合同号", "").strip(),
    }

    # 免租计划 JSON 容错
    try:
        contract_data["免租计划"] = json.loads(request.form.get("免租计划", "[]"))
    except Exception:
        contract_data["免租计划"] = []
    # 保底租金计划 JSON 容错
    try:
        contract_data["保底租金计划"] = json.loads(request.form.get("保底租金计划", "[]"))
    except Exception:
        contract_data["保底租金计划"] = []
    # 提成扣点计划 JSON 容错
    try:
        contract_data["提成扣点计划"] = json.loads(request.form.get("提成扣点计划", "[]"))
    except Exception:
        contract_data["提成扣点计划"] = []
    # 物业费计划 JSON 容错
    try:
        contract_data["物业费计划"] = json.loads(request.form.get("物业费计划", "[]"))
    except Exception:
        contract_data["物业费计划"] = []

    # ── 校验（对齐原型） ──
    today_val = date.today()
    project_filter = auth.get_project_filter()
    all_shops = {s["铺位号"]: s for s in db.load_shops(project_filter)}
    all_contracts = db.load_contracts(project_filter)

    def _err(msg):
        flash(msg, "danger")
        return redirect(url_for("admin_contracts"))

    if is_edit:
        # ── 编辑模式 ──
        phone_val = contract_data["联系电话"]
        if phone_val and not phone_val.isdigit():
            return _err("联系电话只能是数字")
        extra_dep = contract_data["已付补缴押金"]
        try:
            if extra_dep and float(extra_dep) < 0:
                return _err("已付/补缴押金不能为负数")
        except ValueError:
            return _err("已付/补缴押金格式错误")
    else:
        # ── 新增模式 ──
        # 1. 合同号
        c_no = contract_data["合同号"]
        if not c_no:
            return _err("合同号不能为空")
        if any(c["合同号"] == c_no for c in all_contracts):
            return _err(f"合同号 {c_no} 已存在")
        # 2. 商户名称
        if not contract_data["商户名称"]:
            return _err("商户名称不能为空")
        # 3. 经营业态
        if not contract_data["经营业态"]:
            return _err("请选择经营业态")
        # 4. 所属项目
        if not contract_data["所属项目"]:
            return _err("请选择所属项目")
        # 5. 关联铺位
        sn = contract_data["关联铺位号"]
        if not sn:
            return _err("请选择关联铺位号")
        # 6. 铺位是否被占用
        for c in all_contracts:
            if c.get("关联铺位号") == sn and c["合同号"] != c_no:
                return _err(f"铺位 {sn} 已被合同 {c['合同号']} 占用")
        # 7. 保底租金
        rm = contract_data["租金模式"]
        base_rent_plan_arr = contract_data.get("保底租金计划", [])
        comm_plan_arr = contract_data.get("提成扣点计划", [])
        if rm in ("保底", "取高", ""):
            if not base_rent_plan_arr:
                gr = contract_data["保底租金(元/㎡/天)"]
                if not gr or gr == "0":
                    return _err("保底租金不能为空")
                try:
                    gr_float = float(gr)
                    shop = all_shops.get(sn)
                    if shop:
                        base_rent = float(shop.get("基准租金(元/㎡/天)", "0"))
                        if gr_float < base_rent:
                            return _err(f"保底租金不能低于基准租金 {base_rent}")
                except ValueError:
                    return _err("保底租金格式错误")
        # 8. 提成扣点
        commission_val = contract_data["提成租金扣点(%)"]
        if rm in ("提成", "取高"):
            if not comm_plan_arr:
                if not commission_val or commission_val == "0":
                    return _err(f"{rm}模式提成扣点不能为空")
            try:
                cf = float(commission_val)
                if cf <= 0 or cf > 100:
                    return _err("提成扣点需在0-100之间且大于0")
            except ValueError:
                return _err("提成扣点格式错误")
        else:
            if not commission_val:
                commission_val = "0"
            try:
                cf = float(commission_val)
                if cf < 0 or cf > 100:
                    return _err("提成扣点需在0-100之间")
            except ValueError:
                return _err("提成扣点格式错误")
            contract_data["提成租金扣点(%)"] = commission_val
        # 9. 租金模式
        if not rm:
            return _err("请选择租金模式")
        # 10. 签约日期
        sign_str = contract_data["签约日期"]
        if not sign_str:
            return _err("请选择签约日期")
        try:
            sign_dt = datetime.strptime(sign_str, "%Y-%m-%d").date()
        except ValueError:
            return _err("签约日期格式错误")
        # 11. 租赁开始日期
        start_str = contract_data["租赁开始日期"]
        if not start_str:
            return _err("请选择租赁开始日期")
        try:
            start_dt = datetime.strptime(start_str, "%Y-%m-%d").date()
            if start_dt < sign_dt:
                return _err("租赁开始日期不能早于签约日期")
        except ValueError:
            return _err("租赁开始日期格式错误")
        # 12. 租赁结束日期
        end_str = contract_data["租赁结束日期"]
        if not end_str:
            return _err("请选择租赁结束日期")
        try:
            end_dt = datetime.strptime(end_str, "%Y-%m-%d").date()
            if end_dt <= start_dt:
                return _err("租赁结束日期需晚于开始日期")
        except ValueError:
            return _err("租赁结束日期格式错误")
        # 13. 支付周期
        if not contract_data["支付周期"]:
            return _err("请选择支付周期")
        # 14. 押金
        dep_str = contract_data["押金"]
        if dep_str is None or str(dep_str).strip() == "":
            return _err("押金不能为空")
        try:
            dep_float = float(dep_str)
            if dep_float < 0 or dep_float > 999999:
                return _err("押金范围 0~999999")
        except ValueError:
            return _err("押金格式错误")
        # 15. 已付/补缴押金
        extra_dep_str = contract_data.get("已付补缴押金", "0")
        try:
            extra_dep_val = float(extra_dep_str) if extra_dep_str else 0
            if extra_dep_val < 0:
                return _err("已付/补缴押金不能为负数")
        except ValueError:
            return _err("已付/补缴押金格式错误")
        # 16. 联系电话
        phone_val = contract_data["联系电话"]
        if phone_val and not phone_val.isdigit():
            return _err("联系电话只能是数字")

    # 已付/补缴押金 ≤ 押金 - 意向金抵扣押金
    try:
        deposit_val = float(contract_data.get("押金", 0) or 0)
        intent_val = float(contract_data.get("意向金抵扣押金", 0) or 0)
        extra_val = float(contract_data.get("已付补缴押金", 0) or 0)
        if extra_val > deposit_val - intent_val:
            return _err(f"已付/补缴押金不能超过 {max(0, deposit_val - intent_val):.2f}（押金-意向金抵扣押金）")
    except ValueError:
        pass

    # 自动计算合同状态 + 剩余租期
    term_date = contract_data.get("终止日期", "")
    sign_str = contract_data.get("签约日期", "")
    if contract_data.get("终止原因") or term_date:
        contract_data["合同状态"] = "已终止"
        contract_data["剩余租期(天)"] = "0"
    else:
        try:
            if sign_str and datetime.strptime(sign_str, "%Y-%m-%d").date() > today_val:
                contract_data["合同状态"] = "待生效"
            else:
                contract_data["合同状态"] = "执行中"
        except Exception:
            contract_data["合同状态"] = "执行中"
        try:
            end_dt = datetime.strptime(contract_data["租赁结束日期"], "%Y-%m-%d").date()
            remain = max(0, (end_dt - today_val).days)
            contract_data["剩余租期(天)"] = str(remain)
        except Exception:
            contract_data["剩余租期(天)"] = "0"

    db.save_contracts([contract_data])
    op = "修改" if is_edit else "新增"
    db.log_operation(user["用户名"], user["角色"], op, "合同",
                     f"{op}合同 {contract_data['合同号']}", contract_data["合同号"])

    # 自动同步商机阶段为"已转合同"（对齐原型）
    try:
        opps = db.load_opportunities()
        merchant = contract_data.get("商户名称", "").strip().lower()
        updated = False
        # 同步意向金抵扣押金：始终从关联商机同步（商户名称+项目+铺位号三重匹配）
        project = contract_data.get("所属项目", "").strip().lower()
        shop_no = contract_data.get("关联铺位号", "").strip()
        opp_matched = None
        for o in opps:
            om = o.get("商户名称", "").strip().lower()
            op = o.get("意向项目", "").strip().lower()
            os = o.get("意向铺位", "").strip()
            if (om == merchant
                    and op == project
                    and os == shop_no
                    and o.get("意向金去向", "") == "已转押金"
                    and float(o.get("意向金金额(元)", 0) or 0) > 0):
                opp_matched = o
                break
        new_val = opp_matched["意向金金额(元)"] if opp_matched else "0"
        old_val = contract_data.get("意向金抵扣押金", "0")
        if str(new_val) != str(old_val):
            contract_data["意向金抵扣押金"] = new_val
            db.save_contracts([contract_data])
        # 更新商机阶段
        for o in opps:
            if (o.get("商户名称", "").strip().lower() == merchant
                    and o.get("当前阶段", "") != "已转合同"):
                o["当前阶段"] = "已转合同"
                updated = True
        if updated:
            db.save_opportunities(opps)
    except Exception as e:
        print(f"[商机同步] 自动更新商机阶段失败: {e}")

    # 自动同步铺位状态（对齐原型 sync_shop_status）
    try:
        _sync_shop_status()
    except Exception as e:
        print(f"[铺位同步] 自动同步铺位状态失败: {e}")

    flash(f"合同 {contract_data['合同号']} 保存成功", "success")
    return redirect(url_for("admin_contracts"))

@app.route("/admin/contracts/delete/<contract_no>")
@auth.admin_required
def admin_contracts_delete(contract_no):
    """删除合同"""
    user = auth.get_current_user()
    db.delete_contract(contract_no)
    db.log_operation(user["用户名"], user["角色"], "删除", "合同",
                     f"删除合同 {contract_no}", contract_no)
    try:
        _sync_shop_status()
    except Exception:
        pass
    flash(f"合同 {contract_no} 已删除", "success")
    return redirect(url_for("admin_contracts"))

@app.route("/admin/contracts/batch_delete", methods=["POST"])
@auth.admin_required
def admin_contracts_batch_delete():
    """批量删除合同"""
    user = auth.get_current_user()
    contract_nos = request.form.getlist("contract_nos")
    for cno in contract_nos:
        db.delete_contract(cno)
        db.log_operation(user["用户名"], user["角色"], "删除", "合同", f"批量删除合同 {cno}", cno)
    try:
        _sync_shop_status()
    except Exception:
        pass
    flash(f"已删除 {len(contract_nos)} 个合同", "success")
    return redirect(url_for("admin_contracts"))

@app.route("/admin/contracts/terminate", methods=["POST"])
@auth.admin_required
def admin_contracts_terminate():
    """终止合同（设置终止日期 + 终止原因）"""
    user = auth.get_current_user()
    contract_no = request.form.get("contract_no", "").strip()
    term_date = request.form.get("终止日期", "").strip()
    term_reason = request.form.get("终止原因", "").strip()
    if not contract_no:
        flash("合同号不能为空", "danger")
        return redirect(url_for("admin_contracts"))
    if not term_date:
        flash("请选择终止日期", "danger")
        return redirect(url_for("admin_contracts"))
    all_contracts = db.load_contracts()
    found = None
    for c in all_contracts:
        if c["合同号"] == contract_no:
            found = c
            break
    if not found:
        flash(f"合同 {contract_no} 不存在", "danger")
        return redirect(url_for("admin_contracts"))
    found["终止日期"] = term_date
    found["终止原因"] = term_reason
    found["合同状态"] = "已终止"
    found["剩余租期(天)"] = "0"
    db.save_contracts([found])
    db.log_operation(user["用户名"], user["角色"], "终止", "合同",
                     f"终止合同 {contract_no}", contract_no)
    try:
        _sync_shop_status()
    except Exception:
        pass
    flash(f"合同 {contract_no} 已终止", "success")
    return redirect(url_for("admin_contracts"))

# ===================== 管理端 - 租金收缴 =====================
@app.route("/admin/rent")
@auth.admin_required
def admin_rent():
    user = auth.get_current_user()
    project_filter = auth.get_project_filter()
    contracts = db.load_contracts(project_filter)
    shops_cache = db.load_shops(project_filter)
    payments_cache = db.load_payments(project_filter)
    biz_cache = db.load_business_data(project_filter)

    # 对齐原型：排除「待生效」，保留执行中/已到期/已终止等
    contracts = [c for c in contracts if c.get("合同状态") != "待生效"]

    # 预计算所有合同租金数据（对齐原型 render_list 全部逻辑）
    rent_overview = []
    today = date.today()
    for c in contracts:
        try:
            plan = utils.generate_rent_plan(c, shops_cache, payments_cache, biz_cache)
            # 从 plan 计算全部指标
            tr = round(sum(p["应缴金额(元)"] for p in plan), 2)
            pd = round(sum(p["已缴金额(元)"] for p in plan), 2)
            tp = round(sum(float(p.get("应缴物业费", 0)) for p in plan), 2)
            tpp = round(sum(float(p.get("已缴物业费", 0)) for p in plan), 2)
            te = round(tr + tp, 2)           # 总计预期收入
            trm = round(te - pd - tpp, 2)     # 总计剩余收入
            # 逾期金额（租金逾期 + 物业费逾期）
            ar = 0.0
            prop_ar = 0.0
            for p in plan:
                d = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
                if d < today:
                    ar += max(p["应缴金额(元)"] - p["已缴金额(元)"], 0)
                    prop_ar += max(p.get("应缴物业费", 0) - p.get("已缴物业费", 0), 0)
            ar = round(ar + prop_ar, 2)
            # 租金逾期状态（对齐原型6种状态）
            status = c.get("合同状态", "")
            # 未起租检查
            start_str = c.get("租赁开始日期", "")
            not_started = False
            if start_str:
                try:
                    start_dt = datetime.strptime(start_str, "%Y-%m-%d").date()
                    not_started = today < start_dt
                except Exception:
                    pass

            if status == "已到期":
                next_label, next_days, rent_status = "已到期", -1, "已到期"
            elif status == "已终止":
                next_label, next_days, rent_status = "已终止", -1, "已终止"
            elif not_started:
                next_label, next_days, rent_status = "未起租", -1, "未起租"
            else:
                unpaid = [p for p in plan if p["已缴金额(元)"] < p["应缴金额(元)"]]
                if not unpaid:
                    next_label, next_days, rent_status = "已结清", 9999, "已结清"
                else:
                    # 下次缴费剩余天数：从今天往后找最近的未缴期次
                    future_unpaid = [p for p in unpaid if datetime.strptime(p["支付时间"], "%Y-%m-%d").date() >= today]
                    future_unpaid.sort(key=lambda x: x["支付时间"])
                    if future_unpaid:
                        d = datetime.strptime(future_unpaid[0]["支付时间"], "%Y-%m-%d").date()
                        diff = (d - today).days
                        next_label, next_days = str(diff), diff
                    else:
                        next_label, next_days = "-", -1
                    # 租金逾期状态：是否有任何逾期未缴的期次
                    has_overdue = any(
                        datetime.strptime(p["支付时间"], "%Y-%m-%d").date() < today
                        for p in unpaid
                    )
                    rent_status = "已逾期" if has_overdue else "正常"
            # 物业费逾期状态（对齐原型）
            if status in ("已到期", "已终止"):
                prop_status = status
            elif not_started:
                prop_status = "未起租"
            else:
                prop_unpaid = [
                    p for p in plan
                    if float(p.get("应缴物业费", 0)) > 0
                    and float(p.get("已缴物业费", 0)) < float(p.get("应缴物业费", 0))
                ]
                if not prop_unpaid:
                    prop_status = "已结清"
                else:
                    prop_unpaid.sort(key=lambda x: x["支付时间"])
                    prop_overdue = False
                    for p in prop_unpaid:
                        d = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
                        if d < today:
                            prop_overdue = True
                            break
                    prop_status = "已逾期" if prop_overdue else "正常"

            rent_overview.append({
                "合同号": c.get("合同号", ""),
                "商户名称": c.get("商户名称", ""),
                "免租期(天)": int(c.get("免租期(天)", 0) or 0),
                "总租金": tr,
                "总物业费": tp,
                "总计预期收入": te,
                "总计已收租金": pd,
                "总计已收物业费": tpp,
                "总计剩余收入": trm,
                "下次缴费剩余天数": next_label,
                "下次天数": next_days,
                "租金逾期状态": rent_status,
                "物业费逾期状态": prop_status,
                "逾期金额": ar,
            })
        except Exception as e:
            print(f"[rent] 计算 {c.get('合同号','')} 失败: {e}")
            continue

    return render_template("admin/rent.html", user=user, rent_overview=rent_overview)


@app.route("/api/rent-plan/<contract_no>")
@auth.login_required
def api_rent_plan(contract_no):
    """AJAX：获取合同的租金缴纳计划明细 + 合同详情（对齐原型 show_plan）"""
    try:
        project_filter = auth.get_project_filter()
        contracts = db.load_contracts(project_filter)
        contract = None
        for c in contracts:
            if c.get("合同号") == contract_no:
                contract = c
                break
        if not contract:
            return jsonify({"error": "合同不存在"})
        shops_cache = db.load_shops(project_filter)
        payments_cache = db.load_payments(project_filter)
        biz_cache = db.load_business_data(project_filter)
        # 查出铺位面积信息
        shop_area = shop_space_type = shop_building_area = shop_use_area = shop_area_source = ""
        shop_no = contract.get("关联铺位号", "")
        for shop in shops_cache:
            if str(shop.get("铺位号", "")).strip() == str(shop_no).strip():
                shop_space_type = shop.get("空间类型", "")
                shop_building_area = shop.get("建筑面积(㎡)", "")
                shop_use_area = shop.get("使用面积(㎡)", "")
                shop_area = shop.get("计租面积(㎡)", "")
                if not shop_area:
                    shop_area = shop_building_area
                if shop_space_type == "商铺":
                    shop_area_source = "建筑面积"
                elif shop_space_type == "室外场地":
                    shop_area_source = "使用面积"
                break
        plan = utils.generate_rent_plan(contract, shops_cache, payments_cache, biz_cache)

        # ── 查询关联商机意向金抵扣租金（三重匹配） ──
        opp_deposit = 0.0
        try:
            opps = db.load_opportunities(project_filter)
            merchant = contract.get("商户名称", "").strip().lower()
            project = contract.get("所属项目", "").strip().lower()
            shop_no_key = contract.get("关联铺位号", "").strip()
            for o in opps:
                om = o.get("商户名称", "").strip().lower()
                op = o.get("意向项目", "").strip().lower()
                os = o.get("意向铺位", "").strip()
                if (om == merchant and op == project and os == shop_no_key
                        and o.get("意向金去向", "") == "已转租金"):
                    opp_deposit = float(o.get("意向金金额(元)", 0) or 0)
                    break
        except Exception:
            opp_deposit = 0.0

        # 计算汇总（对齐原型）
        tr = round(sum(p["应缴金额(元)"] for p in plan), 2)
        pd = round(sum(p["已缴金额(元)"] for p in plan), 2)
        rm = round(tr - pd, 2)
        ar = 0.0
        today = date.today()
        for p in plan:
            d = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
            if d < today:
                ar += max(p["应缴金额(元)"] - p["已缴金额(元)"], 0)
                ar += max(p.get("应缴物业费", 0) - p.get("已缴物业费", 0), 0)
        ar = round(ar, 2)
        return jsonify({
            "plan": plan,
            "contract": {
                "合同号": contract.get("合同号", ""),
                "商户名称": contract.get("商户名称", ""),
                "经营业态": contract.get("经营业态", ""),
                "所属项目": contract.get("所属项目", ""),
                "关联铺位号": contract.get("关联铺位号", ""),
                "联系人": contract.get("联系人", ""),
                "联系电话": contract.get("联系电话", ""),
                "租金模式": contract.get("租金模式", "保底"),
                "保底租金(元/㎡/天)": contract.get("保底租金(元/㎡/天)", ""),
                "提成租金扣点(%)": contract.get("提成租金扣点(%)", ""),
                "物业服务费单价（元/㎡/天）": contract.get("物业服务费单价（元/㎡/天）", ""),
                "押金": contract.get("押金", ""),
                "押金支付状态": contract.get("押金支付状态", ""),
                "意向金抵扣押金": contract.get("意向金抵扣押金", "0"),
                "已付补缴押金": contract.get("已付补缴押金", "0"),
                "签约日期": contract.get("签约日期", ""),
                "租赁开始日期": contract.get("租赁开始日期", ""),
                "租赁结束日期": contract.get("租赁结束日期", ""),
                "免租期(天)": contract.get("免租期(天)", ""),
                "终止日期": contract.get("终止日期", ""),
                "合同状态": contract.get("合同状态", ""),
                "剩余租期(天)": contract.get("剩余租期(天)", ""),
                "备注": contract.get("备注", ""),
            "建筑面积": shop_area,
            "空间类型": shop_space_type,
            "计租面积来源": shop_area_source,
            "意向金抵扣租金": opp_deposit,
            },
            "summary": {
                "总租金": tr,
                "已收租金": pd,
                "尚欠租金": rm,
                "逾期金额": ar,
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/rent-update-paid", methods=["POST"])
@auth.login_required
def api_rent_update_paid():
    """AJAX：更新缴费金额"""
    data = request.get_json()
    cno = data.get("合同号", "")
    pay_time = data.get("支付时间", "")
    val = data.get("金额", 0)
    val_type = data.get("类型", "租金")  # "租金" 或 "物业费"
    if not cno or not pay_time:
        return jsonify({"ok": False, "error": "参数缺失"})
    try:
        if val_type == "物业费":
            db.update_property_fee_paid(cno, pay_time, float(val))
        else:
            db.update_paid(cno, pay_time, float(val))
        user = auth.get_current_user()
        db.log_operation(user["用户名"], user["角色"], "编辑缴费", "租金收缴",
                         f"合同{cno} {pay_time} {val_type}={val}")
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


def _build_demand_letter_doc(contract):
    """构建单个催缴函 Document 对象（严格参照模板格式）"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    project_filter = auth.get_project_filter()
    shops_cache = db.load_shops(project_filter)
    payments_cache = db.load_payments(project_filter)
    biz_cache = db.load_business_data(project_filter)

    today = date.today()
    ymd = today.strftime("%Y年%m月%d日")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _set_font(run, size=None, bold=None):
        run.font.name = '微软雅黑'
        if size:
            run.font.size = size
        if bold is not None:
            run.bold = bold
        rPr = run._element.get_or_add_rPr()
        existing = rPr.findall(qn('w:rFonts'))
        found = False
        for e in existing:
            if e.get(qn('w:eastAsia')) == '微软雅黑':
                found = True
                break
        if not found:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            rPr.insert(0, rFonts)

    def _set_indent(para, chars=2):
        pPr = para._element.get_or_add_pPr()
        for old in pPr.findall(qn('w:ind')):
            pPr.remove(old)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), str(chars * 280))
        ind.set(qn('w:firstLineChars'), str(chars * 100))
        pPr.append(ind)

    def _set_para_spacing(para):
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

    def _add_body(text, indent_chars=2):
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_font(run, size=Pt(14))
        _set_para_spacing(p)
        if indent_chars > 0:
            _set_indent(p, indent_chars)
        return p

    def _add_empty(indent_chars=0):
        p = doc.add_paragraph()
        _set_para_spacing(p)
        if indent_chars > 0:
            _set_indent(p, indent_chars)
        return p

    # 标题
    p = doc.add_paragraph()
    run = p.add_run("租金催缴函")
    _set_font(run, size=Pt(16), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_empty(indent_chars=2)

    # 称呼
    p = doc.add_paragraph()
    run = p.add_run(f"尊敬的 {contract['商户名称']}：")
    _set_font(run, size=Pt(14))
    _set_para_spacing(p)

    _add_empty(indent_chars=2)

    # ── 欠缴信息：分别计算租金和物业费 ──
    plan = utils.generate_rent_plan(contract, shops_cache, payments_cache, biz_cache)
    rent_rows = []
    prop_rows = []
    rent_sub = 0.0
    prop_sub = 0.0
    for p_data in plan:
        d = datetime.strptime(p_data["支付时间"], "%Y-%m-%d").date()
        if d < today:
            rd = round(max(p_data["应缴金额(元)"] - p_data["已缴金额(元)"], 0), 2)
            ps = round(float(p_data.get("应缴物业费", 0) or 0), 2)
            pp = round(float(p_data.get("已缴物业费", 0) or 0), 2)
            pd = round(max(ps - pp, 0), 2)
            if rd > 0:
                rent_rows.append((p_data["支付时间"], round(p_data["应缴金额(元)"], 2), rd))
                rent_sub += rd
            if pd > 0:
                prop_rows.append((p_data["支付时间"], ps, pd))
                prop_sub += pd
    rent_sub = round(rent_sub, 2)
    prop_sub = round(prop_sub, 2)
    arrears_total = round(rent_sub + prop_sub, 2)

    # 费用合计：直接嵌入正文，不加框不加粗
    _add_body(f"基于贵我双方签订的{contract['合同号']}合同，您已欠缴我司费用合计{arrears_total:.2f}元，详细情况如下：")

    _add_empty()

    p = doc.add_paragraph()
    run = p.add_run(f"合同编号：{contract['合同号']}")
    _set_font(run, size=Pt(10.5))
    _set_para_spacing(p)

    p = doc.add_paragraph()
    run = p.add_run(f"租赁起止日期：{contract['租赁开始日期']} 至 {contract['租赁结束日期']}")
    _set_font(run, size=Pt(10.5))
    _set_para_spacing(p)

    _add_empty()

    # ── 逾期明细表格辅助函数（4列：序号/应支付日期/应缴金额/欠付金额） ──
    def _add_arrears_table(title, headers, data_rows, should_sub, subtotal_val, col_widths):
        p = doc.add_paragraph()
        run = p.add_run(title)
        _set_font(run, size=Pt(11), bold=True)
        _set_para_spacing(p)

        ncols = len(headers)
        tbl = doc.add_table(rows=1, cols=ncols)
        tbl.style = 'Table Grid'
        tw = sum(col_widths)
        _el = tbl._tbl
        _pr = _el.find(qn('w:tblPr'))
        if _pr is None:
            _pr = OxmlElement('w:tblPr')
            _el.insert(0, _pr)
        _tw = _pr.find(qn('w:tblW'))
        if _tw is None:
            _tw = OxmlElement('w:tblW')
            _pr.insert(0, _tw)
        _tw.set(qn('w:w'), str(tw))
        _tw.set(qn('w:type'), 'dxa')
        _gr = _el.find(qn('w:tblGrid'))
        if _gr is None:
            _gr = OxmlElement('w:tblGrid')
            _el.insert(1, _gr)
        for _gc in _gr.findall(qn('w:gridCol')):
            _gr.remove(_gc)
        for w_val in col_widths:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(w_val))
            _gr.append(gc)

        for i, text in enumerate(headers):
            c = tbl.rows[0].cells[i]
            c.text = ""
            hp = c.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = hp.add_run(text)
            _set_font(hr, size=Pt(10.5))
            _set_para_spacing(hp)

        # 数据行（4列：序号/日期/应缴/欠付）
        for seq, rd in enumerate(data_rows, 1):
            r = tbl.add_row().cells
            vals = [str(seq), str(rd[0]), f"{rd[1]:.2f}", f"{rd[2]:.2f}"]
            for i, val in enumerate(vals):
                r[i].text = ""
                vp = r[i].paragraphs[0]
                vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                vr = vp.add_run(val)
                _set_font(vr, size=Pt(10.5))
                _set_para_spacing(vp)

        # 小计行：序号+日期合并为"小计"
        sub_cells = tbl.add_row().cells
        # 合并第1、2列
        sub_cells[0]._tc.get_or_add_tcPr().append(
            _merge_elem(1, 1)
        )
        sub_cells[0].text = ""
        sp0 = sub_cells[0].paragraphs[0]
        sp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr0 = sp0.add_run("小计")
        _set_font(sr0, size=Pt(10.5), bold=True)
        _set_para_spacing(sp0)

        sub_cells[2].text = ""
        sp2 = sub_cells[2].paragraphs[0]
        sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr2 = sp2.add_run("\u2014")
        _set_font(sr2, size=Pt(10.5))
        _set_para_spacing(sp2)

        sub_cells[3].text = ""
        sp3 = sub_cells[3].paragraphs[0]
        sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr3 = sp3.add_run(f"{subtotal_val:.2f}")
        _set_font(sr3, size=Pt(10.5), bold=True)
        _set_para_spacing(sp3)

        for row_obj in tbl.rows:
            for i, cell in enumerate(row_obj.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.insert(0, tcW)
                tcW.set(qn('w:w'), str(col_widths[i]))
                tcW.set(qn('w:type'), 'dxa')

        _add_empty()

    def _merge_elem(vm_start, vm_end):
        """生成 w:vMerge 元素用于水平合并"""
        vm = OxmlElement('w:gridSpan')
        vm.set(qn('w:val'), str(vm_end - vm_start + 1))
        return vm

    # 4列宽度：序号/应支付日期/应缴金额(元)/欠付金额(元)
    CW4 = [1000, 2182, 2690, 2864]

    # 租金欠缴明细表
    if rent_rows:
        rent_should_sub = round(sum(r[1] for r in rent_rows), 2)
        _add_arrears_table("租金欠缴明细：",
                           ["序号", "应支付日期", "应缴金额（元）", "欠付金额（元）"],
                           rent_rows, rent_should_sub, rent_sub, CW4)

    # 物业费欠缴明细表
    if prop_rows:
        prop_should_sub = round(sum(r[1] for r in prop_rows), 2)
        _add_arrears_table("物业费欠缴明细：",
                           ["序号", "应支付日期", "应缴物业费（元）", "欠付物业费（元）"],
                           prop_rows, prop_should_sub, prop_sub, CW4)

    _add_body("请于收函之日起3日内将欠缴费用支付至以下账户，否则我们将采取进一步措施。")

    _add_empty()

    _add_body("开户名称：XXX公司", indent_chars=4)
    _add_body("开户行：XXX银行", indent_chars=4)
    _add_body("账号：11111111111", indent_chars=4)

    _add_empty(indent_chars=2)

    _add_body("顺颂商祺。")

    _add_empty()

    p_company = doc.add_paragraph()
    run = p_company.add_run("xxx公司")
    _set_font(run, size=Pt(14))
    _set_para_spacing(p_company)
    p_company.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_date = doc.add_paragraph()
    run = p_date.add_run(ymd)
    _set_font(run, size=Pt(14))
    _set_para_spacing(p_date)
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


@app.route("/api/generate-demand-letter", methods=["POST"])
@auth.login_required
def api_generate_demand_letter():
    """生成单个催缴函 Word，直接返回 .docx 文件下载"""
    # 兼容 JSON 和 form 两种请求方式
    if request.is_json:
        data = request.get_json()
    elif request.form.get('data'):
        import json
        data = json.loads(request.form['data'])
    else:
        data = {}
    contract_nos = data.get("contract_nos", [])
    if not contract_nos:
        return jsonify({"ok": False, "error": "请先选择合同"})

    cno = contract_nos[0]  # 每次只处理一个合同
    project_filter = auth.get_project_filter()
    contracts = db.load_contracts(project_filter)
    contract = None
    for c in contracts:
        if c.get("合同号") == cno:
            contract = c
            break
    if not contract:
        return jsonify({"ok": False, "error": f"合同 {cno} 不存在"})

    try:
        doc = _build_demand_letter_doc(contract)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"{cno} {contract['商户名称']}.docx"
        # 记录操作日志
        user = auth.get_current_user()
        db.log_operation(user["用户名"], user["角色"], "生成催缴函", "租金收缴", f"合同：{cno}")
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/search-contracts")
@auth.login_required
def api_search_contracts():
    """AJAX：搜索合同（按合同号/商户名）"""
    q = request.args.get("q", "").strip().lower()
    project_filter = auth.get_project_filter()
    contracts = db.load_contracts(project_filter)
    if not q:
        return jsonify([])
    results = []
    for c in contracts:
        if q in (c.get("合同号", "") + c.get("商户名称", "")).lower():
            results.append({
                "合同号": c.get("合同号", ""),
                "商户名称": c.get("商户名称", ""),
                "经营业态": c.get("经营业态", ""),
            })
    return jsonify(results[:10])

# ===================== 管理端 - 经营数据 =====================
@app.route("/admin/business")
@auth.admin_required
def admin_business():
    user = auth.get_current_user()
    project_filter = auth.get_project_filter()
    biz_data = db.load_business_data(project_filter)
    # 按日期降序排列
    biz_data = sorted(biz_data, key=lambda x: x.get("日期", ""), reverse=True)
    from config import PROJECT_OPTIONS, BUSINESS_TYPE
    contracts_data = db.load_contracts(project_filter)
    return render_template("admin/business.html", user=user, biz_data=biz_data,
                           projects=PROJECT_OPTIONS, biz_types=BUSINESS_TYPE,
                           contracts_data=contracts_data)


@app.route("/admin/business/add", methods=["POST"])
@auth.admin_required
def admin_business_add():
    """新增/修改经营数据"""
    user = auth.get_current_user()
    cno = request.form.get("合同号", "").strip()
    biz_date = request.form.get("日期", "").strip()
    revenue = request.form.get("营业额", "0").strip()
    footfall = request.form.get("客流量", "").strip()
    deals = request.form.get("成交量", "").strip()
    remark = request.form.get("备注", "").strip()

    errors = []
    if not cno: errors.append("请搜索选择商户")
    if not biz_date: errors.append("请选择日期")
    try: revenue = float(revenue)
    except: errors.append("营业额格式错误")

    fv = "" if not footfall else int(footfall)
    dv = "" if not deals else int(deals)

    if errors:
        for e in errors: flash(e, "danger")
        return redirect(url_for("admin_business"))

    # 获取商户信息
    contracts = db.load_contracts()
    contract = None
    for c in contracts:
        if c.get("合同号") == cno:
            contract = c
            break
    if not contract:
        flash("未找到对应合同", "danger")
        return redirect(url_for("admin_business"))

    record = {
        "合同号": cno,
        "商户名称": contract.get("商户名称", ""),
        "日期": biz_date,
        "营业额": round(revenue, 2),
        "客流量": fv,
        "成交量": dv,
        "业态": contract.get("经营业态", ""),
        "录入时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "备注": remark,
        "所属项目": contract.get("所属项目", ""),
    }
    db.save_business_data([record])
    db.log_operation(user["用户名"], user["角色"], "新增", "经营数据",
                     f"{cno} {biz_date} ¥{revenue}")
    flash("经营数据保存成功", "success")
    return redirect(url_for("admin_business"))


@app.route("/admin/business/delete", methods=["POST"])
@auth.admin_required
def admin_business_delete():
    """批量删除经营数据"""
    user = auth.get_current_user()
    ids = request.form.getlist("ids")
    if not ids:
        flash("请选择要删除的记录", "warning")
        return redirect(url_for("admin_business"))
    for item in ids:
        parts = item.split("||")
        if len(parts) == 2:
            db.delete_business_data(parts[0], parts[1])
    db.log_operation(user["用户名"], user["角色"], "删除", "经营数据",
                     f"批量删除 {len(ids)} 条")
    flash(f"已删除 {len(ids)} 条记录", "success")
    return redirect(url_for("admin_business"))


@app.route("/admin/business/export")
@auth.admin_required
def admin_business_export():
    """导出经营数据为CSV"""
    project_filter = auth.get_project_filter()
    biz_data = db.load_business_data(project_filter)
    biz_data = sorted(biz_data, key=lambda x: x.get("日期", ""), reverse=True)

    import io, csv
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["序号", "日期", "合同号", "商户名称", "营业额",
                     "客流量", "成交量", "业态", "所属项目", "录入时间", "备注"])
    for i, b in enumerate(biz_data, 1):
        writer.writerow([i, b.get("日期",""), b.get("合同号",""), b.get("商户名称",""),
                        b.get("营业额",""), b.get("客流量",""), b.get("成交量",""),
                        b.get("业态",""), b.get("所属项目",""), b.get("录入时间",""),
                        b.get("备注","")])

    from flask import Response
    output.seek(0)
    return Response(
        output.getvalue().encode("utf-8-sig"),
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment;filename=business_data_{date.today()}.csv"}
    )

# ===================== 管理端 - 商机管理 =====================
OPPORTUNITY_STAGES = ["初步接洽", "需求确认", "谈判协商", "意向确认", "已支付意向金", "已转合同"]
FOLLOW_RESULTS = ["继续跟进", "阶段推进", "已流失", "已转合同"]
DEPOSIT_DESTINATIONS = ["", "已转押金", "已退还", "已转租金"]

@app.route("/admin/opportunities")
@auth.admin_required
def admin_opportunities():
    user = auth.get_current_user()
    project_filter = auth.get_project_filter()
    opps = db.load_opportunities(project_filter)
    shops = db.load_shops(project_filter)
    from config import PROJECT_OPTIONS, BUSINESS_TYPE
    return render_template("admin/opportunities.html", user=user, opps=opps,
                           shops=shops, projects=PROJECT_OPTIONS, biz_types=BUSINESS_TYPE,
                           stages=OPPORTUNITY_STAGES, follow_results=FOLLOW_RESULTS,
                           deposit_dest=DEPOSIT_DESTINATIONS)


@app.route("/admin/opportunities/add", methods=["POST"])
@auth.admin_required
def admin_opportunities_add():
    """新增/编辑商机"""
    user = auth.get_current_user()
    opp_no = request.form.get("商机编号", "").strip()
    merchant = request.form.get("商户名称", "").strip()

    if not opp_no or not merchant:
        flash("商机编号和商户名称为必填", "danger")
        return redirect(url_for("admin_opportunities"))

    # 保留原有数据（编辑时保留意向金、跟进结果等不在表单中的字段）
    existing_opp = None
    all_opps = db.load_opportunities()
    for o in all_opps:
        if o.get("商机编号") == opp_no:
            existing_opp = o
            break

    data = {
        "商机编号": opp_no,
        "商户名称": merchant,
        "联系人": request.form.get("联系人", "").strip(),
        "联系电话": request.form.get("联系电话", "").strip(),
        "意向主体": request.form.get("意向主体", "").strip(),
        "意向项目": request.form.get("意向项目", "").strip(),
        "意向业态": request.form.get("意向业态", "").strip(),
        "意向租赁期限（年）": request.form.get("意向租赁期限", "").strip(),
        "意向铺位": request.form.get("意向铺位", "").strip(),
        "建筑面积(㎡)": request.form.get("建筑面积", "").strip(),
        "意向租金单价(元/㎡/天)": request.form.get("意向租金单价", "").strip(),
        "物业服务费单价(元/㎡/月)": request.form.get("物业服务费单价", "").strip(),
        "支付周期": request.form.get("支付周期", "").strip(),
        "商机来源": request.form.get("商机来源", "").strip(),
        "当前阶段": request.form.get("当前阶段", "").strip() or "初步接洽",
        "首次接洽日期": request.form.get("首次接洽日期", "").strip(),
        "负责人": request.form.get("负责人", "").strip(),
        "备注": request.form.get("备注", "").strip(),
    }

    # 编辑模式：保留意向金、跟进结果、最近跟进日期、跟进记录
    if existing_opp:
        data["意向金金额(元)"] = existing_opp.get("意向金金额(元)", "")
        data["意向金支付日期"] = existing_opp.get("意向金支付日期", "")
        data["意向金去向"] = existing_opp.get("意向金去向", "")
        data["跟进结果"] = existing_opp.get("跟进结果", "")
        data["最近跟进日期"] = existing_opp.get("最近跟进日期", "")
        data["跟进记录"] = existing_opp.get("跟进记录", [])
    else:
        # 新增：默认值
        data["意向金金额(元)"] = ""
        data["意向金支付日期"] = ""
        data["意向金去向"] = ""
        data["跟进结果"] = "继续跟进"
        data["最近跟进日期"] = ""
        data["跟进记录"] = []

    db.save_opportunities([data])
    db.log_operation(user["用户名"], user["角色"],
                     "修改" if existing_opp else "新增", "商机",
                     f"{'修改' if existing_opp else '新增'}商机 {opp_no}", opp_no)
    flash(f"商机 {opp_no} 保存成功", "success")
    return redirect(url_for("admin_opportunities"))


@app.route("/admin/opportunities/delete/<opp_no>")
@auth.admin_required
def admin_opportunities_delete(opp_no):
    """删除商机"""
    user = auth.get_current_user()
    db.delete_opportunity(opp_no)
    db.log_operation(user["用户名"], user["角色"], "删除", "商机",
                     f"删除商机 {opp_no}", opp_no)
    flash(f"商机 {opp_no} 已删除", "success")
    return redirect(url_for("admin_opportunities"))

@app.route("/admin/opportunities/batch_delete", methods=["POST"])
@auth.admin_required
def admin_opportunities_batch_delete():
    """批量删除商机"""
    user = auth.get_current_user()
    opp_nos = request.form.getlist("opp_nos")
    if not opp_nos:
        flash("未选择任何商机", "danger")
        return redirect(url_for("admin_opportunities"))
    for opp_no in opp_nos:
        db.delete_opportunity(opp_no)
        db.log_operation(user["用户名"], user["角色"], "批量删除", "商机",
                         f"批量删除商机 {opp_no}", opp_no)
    flash(f"已删除 {len(opp_nos)} 个商机", "success")
    return redirect(url_for("admin_opportunities"))


@app.route("/api/opportunity/follow-up", methods=["POST"])
@auth.login_required
def api_opportunity_follow_up():
    """AJAX：添加跟进记录（字段名对齐桌面原型）"""
    data = request.get_json()
    opp_no = data.get("商机编号", "")
    record = data.get("记录", {})

    if not opp_no or not record:
        return jsonify({"ok": False, "error": "参数缺失"})

    content = record.get("跟进内容", "")
    if not content:
        return jsonify({"ok": False, "error": "跟进内容不能为空"})

    opps = db.load_opportunities()
    opp = None
    for o in opps:
        if o.get("商机编号") == opp_no:
            opp = o
            break
    if not opp:
        return jsonify({"ok": False, "error": "商机不存在"})

    follow_records = opp.get("跟进记录", [])
    if isinstance(follow_records, str):
        try:
            follow_records = json.loads(follow_records)
        except:
            follow_records = []
    if not isinstance(follow_records, list):
        follow_records = []

    follow_records.append(record)
    opp["跟进记录"] = follow_records
    opp["最近跟进日期"] = record.get("跟进日期", "")
    opp["当前阶段"] = record.get("阶段", opp.get("当前阶段", ""))
    opp["跟进结果"] = record.get("跟进结果", opp.get("跟进结果", ""))

    db.save_opportunities(opps)
    user = auth.get_current_user()
    db.log_operation(user["用户名"], user["角色"], "跟进", "商机",
                     f"商机{opp_no} 跟进", opp_no)
    return jsonify({"ok": True})


@app.route("/api/opportunity/deposit", methods=["POST"])
@auth.login_required
def api_opportunity_deposit():
    """AJAX：支付意向金（对齐桌面原型逻辑）"""
    data = request.get_json()
    opp_no = data.get("商机编号", "")
    amount = data.get("金额", "0")
    pay_date = data.get("支付日期", "")
    destination = data.get("去向", "")
    remark = data.get("备注", "")

    if not opp_no:
        return jsonify({"ok": False, "error": "参数缺失"})
    if not amount:
        return jsonify({"ok": False, "error": "请填写意向金金额"})
    try:
        if float(amount) <= 0:
            return jsonify({"ok": False, "error": "意向金金额需大于0"})
    except:
        return jsonify({"ok": False, "error": "意向金金额格式错误"})

    from datetime import date as _date
    today = _date.today().strftime("%Y-%m-%d")

    opps = db.load_opportunities()
    opp = None
    for o in opps:
        if o.get("商机编号") == opp_no:
            opp = o
            break
    if not opp:
        return jsonify({"ok": False, "error": "商机不存在"})

    opp["意向金金额(元)"] = str(amount)
    opp["意向金支付日期"] = pay_date
    opp["意向金去向"] = destination
    opp["当前阶段"] = "已支付意向金"
    opp["最近跟进日期"] = today
    opp["跟进结果"] = "阶段推进"
    if remark:
        opp["备注"] = remark

    # 自动追加一条跟进记录（与桌面原型一致）
    follow_records = opp.get("跟进记录", [])
    if isinstance(follow_records, str):
        try:
            follow_records = json.loads(follow_records)
        except:
            follow_records = []
    if not isinstance(follow_records, list):
        follow_records = []
    follow_records.append({
        "跟进日期":    today,
        "阶段":        "已支付意向金",
        "跟进内容":    f"支付意向金 ¥{amount} 元",
        "跟进结果":    "阶段推进",
        "下次计划日期": "",
        "跟进人":      opp.get("负责人", ""),
    })
    opp["跟进记录"] = follow_records

    db.save_opportunities(opps)
    user = auth.get_current_user()
    db.log_operation(user["用户名"], user["角色"], "意向金", "商机",
                     f"商机{opp_no} 支付意向金 ¥{amount}", opp_no)
    return jsonify({"ok": True})


@app.route("/api/shops-by-project")
@auth.login_required
def api_shops_by_project():
    """AJAX：按项目和业态筛选铺位，返回完整对象供弹窗表格展示"""
    project = request.args.get("project", "").strip()
    biz_types_raw = request.args.get("biz_types", "").strip()
    shops = db.load_shops()
    # biz_types 可能是多个业态（中文顿号分隔）
    results = []
    for s in shops:
        if s.get("铺位状态", "") != "空置":
            continue
        if project and s.get("所属项目", "") != project:
            continue
        results.append({
            "铺位号": s.get("铺位号", ""),
            "所属项目": s.get("所属项目", ""),
            "建筑面积(㎡)": s.get("计租面积(㎡)", s.get("建筑面积(㎡)", "")),
            "基准租金(元/㎡/天)": s.get("基准租金(元/㎡/天)", ""),
            "铺位状态": s.get("铺位状态", ""),
        })
    return jsonify(results[:50])


@app.route("/api/contract-precheck")
@auth.login_required
def api_contract_precheck():
    """AJAX：合同提交前预检查（合同号重复 + 铺位占用）"""
    result = {"dup_no": False, "dup_shop": False}
    c_no = request.args.get("no", "").strip()
    shop_no = request.args.get("shop", "").strip()
    pf = auth.get_project_filter()
    all_contracts = db.load_contracts(pf)
    if c_no:
        result["dup_no"] = any(c["合同号"] == c_no for c in all_contracts)
    if shop_no:
        for c in all_contracts:
            if c.get("关联铺位号") == shop_no and c["合同号"] != c_no:
                result["dup_shop"] = c["合同号"]
                break
    return jsonify(result)

@app.route("/api/shops-for-contract")
@auth.login_required
def api_shops_for_contract():
    """AJAX：合同模块选择铺位，按项目筛选空置铺位"""
    project = request.args.get("project", "").strip()
    edit_shop = request.args.get("edit_shop", "").strip()
    exclude_no = request.args.get("exclude_no", "").strip()  # 续签时排除自身合同
    shops = db.load_shops()
    contracts = db.load_contracts()
    results = []
    for s in shops:
        sno = s.get("铺位号", "")
        if s.get("铺位状态", "") != "空置":
            continue
        if project and s.get("所属项目", "") != project:
            continue
        results.append({
            "铺位号": sno,
            "所属项目": s.get("所属项目", ""),
            "建筑面积(㎡)": s.get("计租面积(㎡)", s.get("建筑面积(㎡)", "")),
            "基准租金(元/㎡/天)": s.get("基准租金(元/㎡/天)", ""),
            "铺位状态": s.get("铺位状态", ""),
        })
    return jsonify(results[:50])

# ===================== 管理端 - 用户管理 =====================
@app.route("/admin/users")
@auth.group_required
def admin_users():
    user = auth.get_current_user()
    users = db.get_all_users()
    from config import PROJECT_OPTIONS, ROLES
    return render_template("admin/users.html", user=user, users=users,
                           projects=PROJECT_OPTIONS, roles=ROLES)

@app.route("/admin/users/add", methods=["POST"])
@auth.group_required
def admin_users_add():
    user = auth.get_current_user()
    user_data = {
        "用户名": request.form.get("用户名", "").strip(),
        "密码": request.form.get("密码", "").strip(),
        "角色": request.form.get("角色", "").strip(),
        "所属项目": request.form.get("所属项目", "").strip(),
    }
    db.save_user(user_data)
    db.log_operation(user["用户名"], user["角色"], "新增/修改", "用户",
                     f"用户管理 {user_data['用户名']}", user_data["用户名"])
    flash(f"用户 {user_data['用户名']} 保存成功", "success")
    return redirect(url_for("admin_users"))

@app.route("/admin/users/delete/<username>")
@auth.group_required
def admin_users_delete(username):
    user = auth.get_current_user()
    if username == user["用户名"]:
        flash("不能删除自己", "danger")
        return redirect(url_for("admin_users"))
    db.delete_user(username)
    db.log_operation(user["用户名"], user["角色"], "删除", "用户",
                     f"删除用户 {username}", username)
    flash(f"用户 {username} 已删除", "success")
    return redirect(url_for("admin_users"))

# ===================== 管理端 - 操作日志 =====================
@app.route("/admin/logs")
@auth.group_required
def admin_logs():
    user = auth.get_current_user()
    logs = db.load_operation_logs()
    return render_template("admin/logs.html", user=user, logs=logs)


# ===================== 商户门户 =====================

@app.route("/merchant")
def merchant_index():
    """商户门户首页"""
    if auth.is_merchant_logged_in():
        return redirect(url_for("merchant_dashboard"))
    return redirect(url_for("merchant_login_page"))


@app.route("/merchant/login", methods=["GET", "POST"])
def merchant_login_page():
    """商户登录页"""
    if auth.is_merchant_logged_in():
        return redirect(url_for("merchant_dashboard"))

    error = ""
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        pwd = request.form.get("password", "").strip()
        if not username or not pwd:
            error = "请输入用户名和密码"
        else:
            c = auth.merchant_login(username, pwd)
            if c:
                db.log_operation("商户: " + c.get("商户名称", ""), "商户", "登录", "商户门户", "商户门户登录成功")
                return redirect(url_for("merchant_dashboard"))
            else:
                error = "用户名或密码不正确，请重新输入"

    return render_template("merchant/login.html", error=error)


@app.route("/merchant/register", methods=["POST"])
def merchant_register():
    """商户注册"""
    if auth.is_merchant_logged_in():
        return redirect(url_for("merchant_dashboard"))

    error = ""
    success = ""
    username = request.form.get("username", "").strip()
    cno = request.form.get("contract_no", "").strip()
    name = request.form.get("merchant_name", "").strip()
    pwd = request.form.get("password", "").strip()
    confirm = request.form.get("confirm_password", "").strip()

    if not username or not cno or not name or not pwd or not confirm:
        error = "请填写所有字段"
    elif pwd != confirm:
        error = "两次输入的密码不一致"
    elif len(pwd) < 6:
        error = "密码长度不能少于6位"
    else:
        result = auth.register_merchant(username, cno, name, pwd)
        if result == "username_exists":
            error = "该用户名已被注册，请更换"
        elif result == "contract_exists":
            error = "该合同号已注册，请直接登录"
        elif result is None:
            error = "合同号或商户名称不正确，请确认后重试"
        else:
            success = "注册成功！请使用用户名和密码登录"

    return render_template("merchant/login.html", error=error, success=success, show_tab="register")


@app.route("/merchant/reset-password", methods=["POST"])
def merchant_reset_password():
    """商户自助找回密码"""
    if auth.is_merchant_logged_in():
        return redirect(url_for("merchant_dashboard"))

    error = ""
    success = ""
    cno = request.form.get("contract_no", "").strip()
    name = request.form.get("merchant_name", "").strip()
    pwd = request.form.get("password", "").strip()
    confirm = request.form.get("confirm_password", "").strip()

    if not cno or not name or not pwd or not confirm:
        error = "请填写所有字段"
    elif pwd != confirm:
        error = "两次输入的密码不一致"
    elif len(pwd) < 6:
        error = "密码长度不能少于6位"
    else:
        result = auth.reset_merchant_password(cno, name, pwd)
        if result == "not_found":
            error = "合同号或商户名称不正确"
        elif result == "not_registered":
            error = "该合同尚未注册账户，请先注册"
        else:
            success = "密码重置成功！请使用新密码登录"

    return render_template("merchant/login.html", error=error, success=success, show_reset=True)


@app.route("/merchant/logout")
def merchant_logout():
    """商户退出登录"""
    db.log_operation("商户: " + auth.get_merchant_contract().get("商户名称", ""),
                     "商户", "登出", "商户门户", "商户门户登出")
    auth.merchant_logout()
    flash("已退出商户门户", "info")
    return redirect(url_for("merchant_login_page"))


@app.route("/merchant/dashboard")
@auth.merchant_required
def merchant_dashboard():
    """商户主页 —— 概要 + 逾期明细"""
    c = auth.get_merchant_contract()
    if not c:
        return redirect(url_for("merchant_login_page"))

    today = date.today()
    plan = utils.generate_rent_plan(c)
    end_dt = datetime.strptime(c["租赁结束日期"], "%Y-%m-%d").date()
    days_left = (end_dt - today).days
    lease_start = c.get("租赁开始日期", "")

    # ── 逾期金额（含租金 + 物业费）──
    overdue_rent = 0.0
    overdue_prop = 0.0
    overdue_items = []
    for i, p in enumerate(plan, 1):
        pay_dt = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
        if pay_dt < today:
            diff_rent = max(p["应缴金额(元)"] - p["已缴金额(元)"], 0)
            diff_prop = max(p["应缴物业费"] - p.get("已缴物业费", 0), 0)
            if diff_rent > 0 or diff_prop > 0:
                overdue_rent += diff_rent
                overdue_prop += diff_prop
                overdue_items.append({
                    "期次": i,
                    "应缴日期": p["支付时间"],
                    "应缴租金": p["应缴金额(元)"],
                    "已缴租金": p["已缴金额(元)"],
                    "欠缴租金": round(diff_rent, 2),
                    "欠缴物业费": round(diff_prop, 2),
                    "合计欠款": round(diff_rent + diff_prop, 2),
                })
    overdue_total = round(overdue_rent + overdue_prop, 2)

    # ── 下次缴费信息（只取未来未结清的最近一期）──
    next_days = "-"
    next_rent = 0.0
    next_prop = 0.0
    for p in plan:
        pay_dt = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
        if pay_dt <= today:
            continue
        diff = max(p["应缴金额(元)"] - p["已缴金额(元)"], 0) + max(p.get("应缴物业费", 0) - p.get("已缴物业费", 0), 0)
        if diff > 0:
            next_days = str((pay_dt - today).days)
            next_rent = p["应缴金额(元)"]
            next_prop = p.get("应缴物业费", 0)
            break

    return render_template("merchant/dashboard.html",
                           contract=c, days_left=days_left,
                           lease_start=lease_start,
                           next_days=next_days, next_rent=next_rent, next_prop=next_prop,
                           overdue_total=overdue_total, overdue_items=overdue_items)


@app.route("/merchant/contract")
@auth.merchant_required
def merchant_contract():
    """商户合同详情"""
    c = auth.get_merchant_contract()
    if not c:
        return redirect(url_for("merchant_login_page"))

    today = date.today()
    end_dt = datetime.strptime(c["租赁结束日期"], "%Y-%m-%d").date()
    days_left = (end_dt - today).days
    status = c.get("合同状态", "")

    # 跳过内部字段
    skip_fields = {"自动租金计划", "免租计划", "保底租金计划", "提成扣点计划", "物业费计划"}
    kv_items = [(k, v) for k, v in c.items() if k not in skip_fields]

    # 免租计划格式化
    free_plans = []
    free_raw = c.get("免租计划", [])
    if isinstance(free_raw, str):
        try:
            free_raw = json.loads(free_raw)
        except:
            free_raw = []
    if isinstance(free_raw, list):
        for i, fp in enumerate(free_raw, 1):
            try:
                s = datetime.strptime(fp["start"], "%Y-%m-%d").date()
                e = datetime.strptime(fp["end"], "%Y-%m-%d").date()
                d = (e - s).days + 1
            except:
                d = 0
            free_plans.append({"序号": i, "开始日期": fp.get("start", ""),
                               "结束日期": fp.get("end", ""), "天数": d})

    return render_template("merchant/contract.html",
                           contract=c, days_left=days_left, status=status,
                           kv_items=kv_items, free_plans=free_plans)


@app.route("/merchant/rent")
@auth.merchant_required
def merchant_rent():
    """商户租金情况"""
    c = auth.get_merchant_contract()
    if not c:
        return redirect(url_for("merchant_login_page"))

    today = date.today()
    plan = utils.generate_rent_plan(c)

    # ── 汇总统计 ──
    total_rent = round(sum(p["应缴金额(元)"] for p in plan), 2)
    paid_rent = round(sum(p["已缴金额(元)"] for p in plan), 2)
    total_prop = round(sum(p.get("应缴物业费", 0) for p in plan), 2)
    paid_prop = round(sum(p.get("已缴物业费", 0) for p in plan), 2)
    remain_rent = round(total_rent - paid_rent, 2)
    remain_prop = round(total_prop - paid_prop, 2)

    # ── 逾期金额 + 明细 ──
    overdue_total = 0.0
    plan_detail = []
    for i, p in enumerate(plan, 1):
        pay_dt = datetime.strptime(p["支付时间"], "%Y-%m-%d").date()
        rent_val = p["应缴金额(元)"]
        paid_val = p["已缴金额(元)"]
        prop_val = p.get("应缴物业费", 0)
        prop_paid = p.get("已缴物业费", 0)
        diff_rent = round(rent_val - paid_val, 2)
        diff_prop = round(prop_val - prop_paid, 2)

        is_overdue = pay_dt < today and (diff_rent > 0 or diff_prop > 0)
        if is_overdue:
            status_label = "逾期"
            status_class = "badge-red"
            overdue_total += diff_rent + diff_prop
        elif diff_rent <= 0 and diff_prop <= 0:
            status_label = "已结清"
            status_class = "badge-green"
        else:
            status_label = "待缴"
            status_class = "badge-blue"

        plan_detail.append({
            "序号": i,
            "支付时间": p["支付时间"],
            "应缴租金": rent_val,
            "应缴物业费": prop_val,
            "本期应缴": round(rent_val + prop_val, 2),
            "已缴租金": paid_val,
            "已缴物业费": prop_paid,
            "状态": status_label,
            "状态CSS": status_class,
            "是否逾期": is_overdue,
        })
    overdue_total = round(overdue_total, 2)

    return render_template("merchant/rent.html",
                           contract=c,
                           total_rent=total_rent, paid_rent=paid_rent,
                           remain_rent=remain_rent,
                           total_prop=total_prop, paid_prop=paid_prop,
                           remain_prop=remain_prop,
                           overdue_total=overdue_total,
                           plan_detail=plan_detail)


@app.route("/merchant/business", methods=["GET", "POST"])
@auth.merchant_required
def merchant_business():
    """商户经营数据 — 录入 + 历史查看 + 时间筛选"""
    c = auth.get_merchant_contract()
    if not c:
        return redirect(url_for("merchant_login_page"))

    cno = c.get("合同号", "")
    name = c.get("商户名称", "")
    today_str = date.today().strftime("%Y-%m-%d")

    # ── POST：内联录入 ──
    if request.method == "POST":
        biz_date = request.form.get("biz_date", today_str).strip()
        revenue = request.form.get("revenue", "").strip()
        footfall = request.form.get("footfall", "").strip()
        deals = request.form.get("deals", "").strip()
        remark = request.form.get("remark", "").strip()

        errors = []
        if not biz_date:
            errors.append("请选择日期")
        if not revenue:
            errors.append("营业额不能为空")
        else:
            try:
                revenue = float(revenue)
            except:
                errors.append("营业额格式错误")
        if footfall:
            try:
                footfall = int(footfall)
            except:
                errors.append("客流量请填写整数")
        else:
            footfall = ""
        deal_val = ""
        if deals:
            try:
                deal_val = int(deals)
            except:
                errors.append("成交量请填写整数")
        else:
            deal_val = ""

        if not errors:
            biz_format = c.get("经营业态", "")
            project = c.get("所属项目", "")
            new_record = {
                "合同号": cno,
                "商户名称": name,
                "日期": biz_date,
                "营业额": round(float(revenue), 2),
                "客流量": footfall,
                "成交量": deal_val,
                "业态": biz_format,
                "录入时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "备注": remark,
                "所属项目": project,
            }
            db.save_business_data([new_record])
            flash("数据保存成功！", "success")
            db.log_operation("商户: " + name, "商户", "录入", "经营数据",
                             f"录入 {biz_date} 营业额 {revenue}")
        else:
            for e in errors:
                flash(e, "danger")
        return redirect(url_for("merchant_business"))

    # ── GET：展示历史 + 录入表单 ──
    start_date = request.args.get("start_date", "").strip()
    end_date = request.args.get("end_date", "").strip()

    all_records = db.load_business_data(contract_filter=cno)
    all_records = sorted(all_records, key=lambda x: x.get("日期", ""), reverse=True)

    if start_date and end_date:
        records = [r for r in all_records if start_date <= r.get("日期", "") <= end_date]
    elif start_date:
        records = [r for r in all_records if r.get("日期", "") >= start_date]
    elif end_date:
        records = [r for r in all_records if r.get("日期", "") <= end_date]
    else:
        records = all_records

    # 统计
    revenues = [r.get("营业额", 0) for r in records if r.get("营业额", 0)]
    deals_list = [int(r.get("成交量", 0)) for r in records if r.get("成交量", "") != ""]
    total_rev = sum(revenues)
    total_deals = sum(deals_list)
    avg_rev = round(total_rev / len(revenues), 2) if revenues else 0
    max_rev = max(revenues) if revenues else 0
    total_footfall = sum(int(r.get("客流量", 0)) for r in records if r.get("客流量", "") != "")
    conv_rate = round(total_deals / total_footfall * 100, 1) if deals_list and total_footfall else 0

    # 今日是否已有录入
    existing_today = any(r.get("日期") == today_str for r in all_records)

    return render_template("merchant/business.html",
                           contract=c, records=records, all_records=all_records,
                           today_str=today_str, existing_today=existing_today,
                           start_date=start_date, end_date=end_date,
                           total_rev=total_rev, avg_rev=avg_rev, max_rev=max_rev,
                           total_footfall=total_footfall, total_deals=total_deals,
                           conv_rate=conv_rate, revenues_count=len(revenues))


# ===================== 启动 =====================
if __name__ == "__main__":
    print("=" * 50)
    print("  LCP商管系统（Web版）")
    print("  管理端: http://localhost:5050/login")
    print("  商户门户: http://localhost:5050/merchant")
    print("=" * 50)
    app.run(host="0.0.0.0", port=5050, debug=True, use_reloader=False)
